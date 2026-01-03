"""Document parsing utilities for Beacon."""

import contextlib
import json
from pathlib import Path

from beacon.models import Document, Query


def load_document(path: Path) -> Document:
    """Load a document from file."""
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    suffix = path.suffix.lower()

    if suffix == ".txt":
        content = _load_text(path)
    elif suffix == ".pdf":
        content = _load_pdf(path)
    elif suffix in (".doc", ".docx"):
        content = _load_docx(path)
    elif suffix == ".md":
        content = _load_text(path)
    elif suffix == ".json":
        content = _load_json_content(path)
    else:
        # Try to read as text
        content = _load_text(path)

    return Document(
        doc_id=path.stem,
        content=content,
        source_path=path,
        metadata={"file_type": suffix, "file_size": path.stat().st_size},
    )


def _load_text(path: Path) -> str:
    """Load plain text file."""
    with open(path, encoding="utf-8") as f:
        return f.read()


def _load_pdf(path: Path) -> str:
    """Load PDF file."""
    try:
        from pypdf import PdfReader
    except ImportError as err:
        raise ImportError(
            "pypdf is required for PDF support. Install with: pip install pypdf"
        ) from err

    reader = PdfReader(path)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return "\n\n".join(text_parts)


def _load_docx(path: Path) -> str:
    """Load DOCX file."""
    try:
        from docx import Document as DocxDocument
    except ImportError as err:
        raise ImportError(
            "python-docx is required for DOCX support. Install with: pip install python-docx"
        ) from err

    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _load_json_content(path: Path) -> str:
    """Load content from JSON file."""
    with open(path, encoding="utf-8") as f:
        data = json.load(f)

    # Try common content field names
    for field in ["content", "text", "body", "document"]:
        if field in data:
            return data[field]  # type: ignore[no-any-return]

    # If it's a string, return it
    if isinstance(data, str):
        return data

    # Otherwise, serialize the whole thing
    return json.dumps(data, indent=2)


def load_documents(paths: list[Path]) -> list[Document]:
    """Load multiple documents from file paths."""
    documents = []
    for path in paths:
        if path.is_dir():
            # Load all supported files from directory
            for file_path in path.iterdir():
                if file_path.is_file() and file_path.suffix.lower() in (
                    ".txt",
                    ".pdf",
                    ".docx",
                    ".md",
                    ".json",
                ):
                    with contextlib.suppress(Exception):
                        documents.append(load_document(file_path))
        else:
            documents.append(load_document(path))
    return documents


def load_queries(path: Path) -> list[Query]:
    """Load queries from JSONL file."""
    if not path.exists():
        raise FileNotFoundError(f"Queries file not found: {path}")

    queries = []
    with open(path, encoding="utf-8") as f:
        for line_num, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                data = json.loads(line)
                query = Query(
                    query_id=data.get("query_id", f"q{line_num}"),
                    text=data["query"],
                    relevant_doc_ids=data.get("relevant_doc_ids", []),
                    relevant_chunk_ids=data.get("relevant_chunk_ids", []),
                    metadata=data.get("metadata", {}),
                )
                queries.append(query)
            except (json.JSONDecodeError, KeyError) as e:
                raise ValueError(f"Error parsing query on line {line_num}: {e}") from e

    return queries


def create_sample_queries(output_path: Path, doc_ids: list[str] | None = None) -> None:
    """Create a sample queries file."""
    sample_queries = [
        {
            "query_id": "q1",
            "query": "What is the main topic of this document?",
            "relevant_doc_ids": doc_ids[:1] if doc_ids else ["doc1"],
        },
        {
            "query_id": "q2",
            "query": "Explain the key concepts discussed.",
            "relevant_doc_ids": doc_ids[:2] if doc_ids and len(doc_ids) >= 2 else ["doc1", "doc2"],
        },
        {
            "query_id": "q3",
            "query": "What are the conclusions or recommendations?",
            "relevant_doc_ids": doc_ids[:1] if doc_ids else ["doc1"],
        },
    ]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        for query in sample_queries:
            f.write(json.dumps(query) + "\n")
