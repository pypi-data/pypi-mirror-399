"""
Word template engine implemented with python-docx + lxml helpers.
Features:
- Text placeholders {{data.field}}
- Table row loops {{#data.list}} with [field] inside row
- Conditional blocks {{?cond}}...{{/}}
- Image placeholders {{@var}}
"""
import tempfile
import shutil
import re
from typing import Any, Dict
from docx import Document

# 作为可安装的 pip 包时，必须使用包内的相对导入
from .utils.text_processor import (
    process_paragraph,
    process_paragraph_blocks,
    process_paragraph_loops,
    _paragraph_has_non_text_elements,
)
from .utils.table_processor import process_table
from .utils.image_processor import process_paragraph_images
from .utils.common import evaluate_condition


def _process_table_conditional_blocks(doc: Document, data: Dict[str, Any]) -> None:
    """
    处理包含表格的条件块。
    格式：{{?condition}}...表格...{{/}}
    如果条件为假，删除条件块内的所有内容（包括表格和段落）。
    """
    start_pat = re.compile(r"\{\{\?([^}]+)\}\}")
    end_pat = re.compile(r"\{\{/\}\}")
    
    # 获取文档 body 的所有子元素（段落和表格）
    body = doc.element.body
    body_elements = list(body)
    
    i = 0
    while i < len(body_elements):
        element = body_elements[i]
        
        # 只处理段落元素
        if element.tag.endswith('}p'):  # paragraph element
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if not para:
                i += 1
                continue
            
            # 获取段落文本
            text = "".join(run.text or "" for run in para.runs)
            start_match = start_pat.search(text)
            
            if not start_match:
                i += 1
                continue
            
            # 找到条件开始标记
            cond = start_match.group(1).strip()
            start_idx = i
            
            # 查找对应的结束标记
            depth = 1
            j = i + 1
            end_idx = None
            
            while j < len(body_elements):
                elem = body_elements[j]
                
                # 检查段落中的条件标记
                if elem.tag.endswith('}p'):
                    para_elem = None
                    for p in doc.paragraphs:
                        if p._element == elem:
                            para_elem = p
                            break
                    if para_elem:
                        t = "".join(run.text or "" for run in para_elem.runs)
                        if start_pat.search(t):
                            depth += 1
                        if end_pat.search(t):
                            depth -= 1
                            if depth == 0:
                                end_idx = j
                                break
                
                j += 1
            
            if end_idx is None:
                # 没有找到匹配的结束标记，跳过
                i += 1
                continue
            
            # 检查条件块内是否包含表格
            # 表格元素的标签格式：{namespace}tbl
            has_table = False
            for k in range(start_idx + 1, end_idx):
                elem_tag = body_elements[k].tag
                # 检查是否是表格元素：标签以 }tbl 结尾或包含 /tbl
                if elem_tag.endswith('}tbl') or elem_tag.endswith('/tbl') or elem_tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl':
                    has_table = True
                    break
            
            # 如果条件块内包含表格，处理条件判断
            if has_table:
                condition_result = evaluate_condition(cond, data)
                
                if not condition_result:
                    # 条件为假：删除条件块内的所有内容（包括表格和段落）
                    # 从后往前删除，避免索引问题
                    elements_to_remove = []
                    for k in range(end_idx, start_idx - 1, -1):
                        elements_to_remove.append(body_elements[k])
                    
                    for elem in elements_to_remove:
                        body.remove(elem)
                    
                    # 重新获取 body_elements（因为删除了元素）
                    body_elements = list(body)
                    # 继续从 start_idx 位置检查（因为删除了元素，索引已经变化）
                    i = max(0, start_idx - 1)
                else:
                    # 条件为真：移除条件标记，保留内容
                    # 移除开始标记
                    new_start_text = start_pat.sub("", text)
                    from .utils.text_processor import _rewrite_paragraph_text
                    
                    if new_start_text.strip() == "":
                        # 如果段落只剩下标记，检查是否包含非文本元素
                        if not _paragraph_has_non_text_elements(para):
                            # 不包含非文本元素，可以安全删除
                            body.remove(element)
                            # 重新获取 body_elements（因为删除了元素）
                            body_elements = list(body)
                            # end_idx 需要减1，因为删除了一个元素
                            if end_idx > start_idx:
                                end_idx -= 1
                        else:
                            # 包含非文本元素，只移除标记文本，保留段落
                            _rewrite_paragraph_text(para, new_start_text)
                    else:
                        # 重写段落文本
                        _rewrite_paragraph_text(para, new_start_text)
                    
                    # 移除结束标记（需要重新查找，因为可能删除了开始段落）
                    if end_idx is not None and end_idx < len(body_elements):
                        end_elem = body_elements[end_idx]
                        if end_elem.tag.endswith('}p'):
                            end_para = None
                            for p in doc.paragraphs:
                                if p._element == end_elem:
                                    end_para = p
                                    break
                            if end_para:
                                end_text = "".join(run.text or "" for run in end_para.runs)
                                new_end_text = end_pat.sub("", end_text)
                                if new_end_text.strip() == "":
                                    # 检查是否包含非文本元素
                                    if not _paragraph_has_non_text_elements(end_para):
                                        # 不包含非文本元素，可以安全删除
                                        body.remove(end_elem)
                                        # 重新获取 body_elements
                                        body_elements = list(body)
                                    else:
                                        # 包含非文本元素，只移除标记文本，保留段落
                                        _rewrite_paragraph_text(end_para, new_end_text)
                                else:
                                    _rewrite_paragraph_text(end_para, new_end_text)
                    
                    # 重新获取 body_elements（确保最新）
                    body_elements = list(body)
                    i += 1
            else:
                # 条件块内没有表格，让 process_paragraph_blocks 处理
                i += 1
        else:
            i += 1


class WordTemplateEngine:
    def __init__(self, template_path: str):
        self.template_path = template_path

    def render(self, data: Dict[str, Any], output_path: str) -> None:
        # work in temp copy to avoid mutating original
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            shutil.copyfile(self.template_path, tmp.name)
            tmp_path = tmp.name

        doc = Document(tmp_path)

        # 首先处理包含表格的条件块（在段落和表格处理之前）
        _process_table_conditional_blocks(doc, data)

        # process tables (row loops + inner placeholders)
        for table in doc.tables:
            process_table(table, data)
            for row in table.rows:
                for cell in row.cells:
                    # handle paragraph-level loops and blocks in cells
                    process_paragraph_loops(cell.paragraphs, data)
                    process_paragraph_blocks(cell.paragraphs, data)
                    for para in cell.paragraphs:
                        process_paragraph(para, data)
                        process_paragraph_images(para, data)

        # process body paragraphs: loops -> blocks -> placeholders/images
        process_paragraph_loops(doc.paragraphs, data)
        process_paragraph_blocks(doc.paragraphs, data)
        for para in doc.paragraphs:
            process_paragraph(para, data)
            process_paragraph_images(para, data)

        doc.save(output_path)


def main():
    sample_data = {
        "data": {
            "diagnose": "测试诊断结果",
            "mutationSiteZero": [
                {"gene_transcript": "BRCA1", "value": "c.1234A>G"},
                {"gene_transcript": "TP53", "value": "c.5678C>T"},
            ],
            "conclusion_A": "",
            "image_config": {
                "url": "123.png",
                
            },
            "detectionResult":{
                'data1':[
                {
                "gene":"gene1",
                'rate':'rate1',
                'age':'age1',
                'symptoms':'symptoms1'
            },
            {
                "gene":"gene2",
                'rate':'rate1',
                'age':'age2',
                'symptoms':'symptoms1'
            }
            ],'data2':[
                {
                "gene":"gene3",
                'rate':'rate3',
                'age':'age3',
                'symptoms':'symptoms3'
            }
            ]
            }
        }
    }
    engine = WordTemplateEngine("template.docx")
    engine.render(sample_data, "output.docx")
    print("done")


if __name__ == "__main__":
    main()

