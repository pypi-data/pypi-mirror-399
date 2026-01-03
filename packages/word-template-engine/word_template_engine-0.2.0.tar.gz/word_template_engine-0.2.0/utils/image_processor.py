"""
Image placeholder processing: {{@var}}
"""
import os
import re
from typing import Any, Dict
from io import BytesIO
import requests
from docx.shared import Inches, Cm
from docx.text.paragraph import Paragraph

from .common import get_nested_value


IMAGE_PATTERN = re.compile(r"\{\{@([\s\S]+?)\}\}", re.DOTALL)


def _parse_size(value: str):
    """Parse size string into a python-docx length (Inches/Cm)."""
    if not value:
        return None
    value = value.lower().strip()
    if value.endswith("px"):
        num = float(value[:-2])
        return Inches(num / 96.0)  # assume 96 DPI
    if value.endswith("cm"):
        num = float(value[:-2])
        return Cm(num)
    if value.endswith("mm"):
        num = float(value[:-2])
        return Cm(num / 10.0)
    if value.endswith("in") or value.endswith("inch"):
        num = float(value[:-2]) if value.endswith("in") else float(value[:-4])
        return Inches(num)
    # no unit: treat as inches
    try:
        num = float(value)
        return Inches(num)
    except Exception:
        return None


def _fetch_image(url: str) -> BytesIO:
    if url.startswith("http://") or url.startswith("https://"):
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        return BytesIO(resp.content)
    with open(url, "rb") as f:
        return BytesIO(f.read())


def process_paragraph_images(paragraph: Paragraph, data: Dict[str, Any]) -> None:
    """Replace {{@var}} with image, tolerant of line breaks across runs, and append image name."""
    full_text = "".join(run.text or "" for run in paragraph.runs)
    if not full_text:
        return
    matches = list(IMAGE_PATTERN.finditer(full_text))
    if not matches:
        return

    cleaned_text = full_text
    image_configs = []

    for match in matches:
        placeholder = match.group(0)
        raw_path = match.group(1)
        var_path = "".join(raw_path.split())
        cfg = get_nested_value(data, var_path)
        image_configs.append((placeholder, cfg))

    # Remove placeholders from text first (to keep layout consistent)
    for placeholder, cfg in image_configs:
        cleaned_text = cleaned_text.replace(placeholder, "", 1)
    print(image_configs)
    # Rewrite paragraph text without placeholders
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = cleaned_text
    else:
        paragraph.add_run(cleaned_text)

    # Append images (and their file names) at the end in order of appearance
    for placeholder, cfg in image_configs:
        if not isinstance(cfg, dict):
            continue
        url = cfg.get("url", "")
        conf = cfg.get("conf", {}) or {}

        width = _parse_size(conf.get("w", "12cm"))
        height = _parse_size(conf.get("h", "8cm"))
        print(url,conf,width,height)
        try:
            img_stream = _fetch_image(url)
            pic_run = paragraph.add_run()
            pic_run.add_picture(img_stream, width=width, height=height)
            # append image file name after picture
        except Exception as exc:
            print(exc)
            error_run = paragraph.add_run(f"[Image Error: {exc}]")

