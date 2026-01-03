"""
Table row loop processing.
"""
import re
import copy
import re
from typing import Any, Dict, List
from docx.table import _Row, Table

from .common import get_nested_value
from .text_processor import render_text, format_bracket_with_item


def _rewrite_paragraph_text(paragraph, new_text: str) -> None:
    """Rewrite paragraph runs sequentially to preserve styles as much as possible."""
    runs = paragraph.runs
    remaining = new_text
    for run in runs:
        if not remaining:
            run.text = ""
            continue
        run_len = len(run.text or "")
        if run_len == 0:
            run.text = remaining
            remaining = ""
        else:
            run.text = remaining[:run_len]
            remaining = remaining[run_len:]
    if remaining and runs:
        # append to last run
        runs[-1].text = (runs[-1].text or "") + remaining


def _strip_loop_marker_in_cell(cell) -> None:
    """Remove {{#...}} marker from a cell while keeping existing run styles (handles split runs)."""
    pattern = re.compile(r"\{\{#([^}]+)\}\}")
    for p in cell.paragraphs:
        full_text = "".join(r.text or "" for r in p.runs)
        new_text = pattern.sub("", full_text)
        if new_text != full_text:
            _rewrite_paragraph_text(p, new_text)


def _clone_row(table: Table, row_idx: int):
    """Clone a row preserving styling."""
    template_row = table.rows[row_idx]
    new_row = template_row._element.clone()
    table._tbl.insert(row_idx + 1, new_row)
    return table.rows[row_idx + 1]


def process_table(table: Table, data: Dict[str, Any]) -> None:
    """
    表格行循环处理：
    1. 表头：{{#mydata}} - 遍历 input_data['mydata']
       - 如果 mydata 是 list: input_data['mydata'] = [{...}, {...}]
       - 如果 mydata 是 dict: input_data['mydata'] = {'data1': [...], 'data2': [...]}
    
    2. 表体行：
       - 如果第一个单元格有 {{+data1}}[name]，第二列是 [age]，找到这行中的 {{/data1}}
         数据来源：input_data['mydata']['data1'] = [{name:1, age:2}, ...]
       - 如果第一个单元格没有 {{+}}，数据来源：input_data['mydata'] = [{name:1, age:2}, ...]
    """
    rows = list(table.rows)
    if not rows:
        return

    # 1) 解析表头 {{#mydata}}
    loop_match = None
    for cell in rows[0].cells:
        m = re.search(r"\{\{#([^}]+)\}\}", cell.text)
        if m:
            loop_match = m
            break
    if not loop_match:
        return

    field_path = loop_match.group(1).strip()
    main_obj = get_nested_value(data, field_path)

    # 去掉表头标记
    for cell in rows[0].cells:
        _strip_loop_marker_in_cell(cell)

    # 2) 处理表体行：扫描每一行，识别是否有 {{+xxx}}...{{/xxx}}
    nested_start_pat = re.compile(r"\{\{\+([^}]+)\}\}")
    nested_end_pat = re.compile(r"\{\{/([^}]+)\}\}")

    # 保存每行的模板和嵌套字段信息
    body_templates: List[tuple[Any, str | None]] = []
    for i in range(1, len(rows)):
        row = rows[i]
        first_cell_text = row.cells[0].text if row.cells else ""
        nested_match = nested_start_pat.search(first_cell_text)
        nested_field = nested_match.group(1).strip() if nested_match else None
        body_templates.append((copy.deepcopy(row._element), nested_field))

    # 3) 删除所有表体行
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    # 4) 处理 main_obj：可能是 list 或 dict
    if isinstance(main_obj, dict):
        # detectionResult 这种情况：{'data1': [...], 'data2': [...]}
        # 遍历每一行模板，根据 {{+xxx}} 从 main_obj 中获取对应的数据列表
        for template_el, nested_field in body_templates:
            if nested_field:
                # 这行有 {{+nested_field}}，从 main_obj[nested_field] 获取数据
                nested_items = main_obj.get(nested_field)
                if not isinstance(nested_items, list):
                    nested_items = []
                
                for nested_item in nested_items:
                    new_el = copy.deepcopy(template_el)
                    table._tbl.append(new_el)
                    new_row = table.rows[-1]
                    
                    for cell in new_row.cells:
                        cell_text = cell.text
                        # 移除 {{+nested_field}} 和 {{/nested_field}} 标记
                        escaped = re.escape(nested_field)
                        cell_text = re.sub(rf"{{{{\+{escaped}}}}}", "", cell_text)
                        cell_text = re.sub(rf"{{{{/{escaped}}}}}", "", cell_text)
                        
                        # 替换 [field] 占位符
                        if isinstance(nested_item, dict):
                            for k, v in nested_item.items():
                                placeholder = f"[{k}]"
                                cell_text = cell_text.replace(placeholder, "" if v is None else str(v))

                        # 处理 [sample|y-m-d] 这类格式，占用 nested_item 作为当前行上下文
                        cell_text = format_bracket_with_item('tab',cell_text, nested_item, data)
                        cell_text = render_text(cell_text, data)
                        
                        # 更新单元格文本
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = ""
                            if p.runs:
                                p.runs[0].text = cell_text
                            else:
                                p.add_run(cell_text)
            else:
                # 没有 {{+}} 的行，跳过（因为 main_obj 是 dict，没有直接数据）
                pass
    elif isinstance(main_obj, list):
        # 普通 list 情况：input_data['mydata'] = [{...}, {...}]
        for main_item in main_obj:
            for template_el, nested_field in body_templates:
                if nested_field:
                    # 嵌套循环：从 main_item[nested_field] 获取数据
                    nested_items = main_item.get(nested_field) if isinstance(main_item, dict) else None
                    if not isinstance(nested_items, list):
                        nested_items = []
                    
                    for nested_item in nested_items:
                        new_el = copy.deepcopy(template_el)
                        table._tbl.append(new_el)
                        new_row = table.rows[-1]
                        
                        for cell in new_row.cells:
                            cell_text = cell.text
                            # 移除 {{+nested_field}} 和 {{/nested_field}} 标记
                            escaped = re.escape(nested_field)
                            cell_text = re.sub(rf"{{{{\+{escaped}}}}}", "", cell_text)
                            cell_text = re.sub(rf"{{{{/{escaped}}}}}", "", cell_text)
                            # cell_text = re.sub(rf"\[\+({escaped})\]", "", cell_text)
                            # cell_text = re.sub(rf"\[/({escaped})\]", "", cell_text)
                            # 替换 [field] 占位符（优先使用 nested_item）
                            if isinstance(nested_item, dict):
                                for k, v in nested_item.items():
                                    placeholder = f"[{k}]"
                                    cell_text = cell_text.replace(placeholder, "" if v is None else str(v))

                            # 如果 nested_item 中没有，尝试从 main_item 获取
                            if isinstance(main_item, dict):
                                for k, v in main_item.items():
                                    placeholder = f"[{k}]"
                                    if placeholder in cell_text:
                                        cell_text = cell_text.replace(placeholder, "" if v is None else str(v))

                            # 处理 [sample|y-m-d]，优先使用 nested_item/main_item 中的字段
                            # 此处传 nested_item，format_bracket_with_item 内部会 fallback 到全局 data
                            cell_text = format_bracket_with_item('tab',cell_text, nested_item, data)
                            cell_text = render_text(cell_text, data)
                            
                            # 更新单元格文本
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.text = ""
                                if p.runs:
                                    p.runs[0].text = cell_text
                                else:
                                    p.add_run(cell_text)
                else:
                    # 没有 {{+}} 的行，直接使用 main_item
                    new_el = copy.deepcopy(template_el)
                    table._tbl.append(new_el)
                    new_row = table.rows[-1]
                    
                    for cell in new_row.cells:
                        cell_text = cell.text
                        
                        # 替换 [field] 占位符
                        if isinstance(main_item, dict):
                            for k, v in main_item.items():
                                placeholder = f"[{k}]"
                                cell_text = cell_text.replace(placeholder, "" if v is None else str(v))

                        # 处理 [sample|y-m-d] 这类格式，使用 main_item 作为当前行上下文
                        cell_text = format_bracket_with_item('tab',cell_text, main_item, data)
                        cell_text = render_text(cell_text, data)
                        
                        # 更新单元格文本
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = ""
                            if p.runs:
                                p.runs[0].text = cell_text
                            else:
                                p.add_run(cell_text)
    else:
        # main_obj 既不是 list 也不是 dict，删除所有表体行
        pass

