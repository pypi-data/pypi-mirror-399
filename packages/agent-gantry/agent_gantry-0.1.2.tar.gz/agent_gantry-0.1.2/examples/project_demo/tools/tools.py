import base64
import datetime
import hashlib
import io
import json
import math
import os
import pathlib
import platform
import random
import re
import shutil
import socket
import statistics
import subprocess
import sys
import tempfile
import time
import uuid
from typing import Any

import pint
import pubchempy as pcp
import requests
from rdkit import Chem
from rdkit.Chem import Descriptors
from sympy import solve, symbols, sympify

from agent_gantry import create_default_gantry

# Create the instance using the factory function
# Users can also call create_default_gantry() for a fresh instance
tools = create_default_gantry()

@tools.register(tags=["chemistry", "molecular"])
def get_molecular_weight(smiles: str) -> float:
    """Calculate the molecular weight of a compound given its SMILES representation."""
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        raise ValueError("Invalid SMILES string.")
    return Descriptors.MolWt(mol)

@tools.register(tags=["chemistry", "compound"])
def get_compound_info(name: str) -> dict[str, Any]:
    """Fetch compound information from PubChem given its name."""
    compounds = pcp.get_compounds(name, 'name')
    if not compounds:
        raise ValueError(f"No compound found for name: {name}")
    compound = compounds[0]
    return {
        "molecular_formula": compound.molecular_formula,
        "molecular_weight": compound.molecular_weight,
        "iupac_name": compound.iupac_name,
        "synonyms": compound.synonyms,
    }

@tools.register(tags=["unit_conversion"])
def convert_units(value: float, from_unit: str, to_unit: str) -> float:
    """Convert a value from one unit to another."""
    ureg = pint.UnitRegistry()
    quantity = value * ureg(from_unit)
    converted = quantity.to(to_unit)
    return converted.magnitude

@tools.register(tags=["date_calculation"])
def calculate_date_difference(date1: str, date2: str) -> int:
    """Calculate the difference in days between two dates (YYYY-MM-DD)."""
    d1 = datetime.datetime.strptime(date1, "%Y-%m-%d")
    d2 = datetime.datetime.strptime(date2, "%Y-%m-%d")
    return abs((d2 - d1).days)

@tools.register(tags=["math", "algebra"])
def solve_equation(equation: str, variable: str) -> Any:
    """Solve a simple algebraic equation for the given variable."""
    var = symbols(variable)
    expr = sympify(equation)
    solution = solve(expr, var)
    return solution

@tools.register(tags=["web"])
def fetch_web_content(url: str) -> str:
    """Fetch the content of a web page given its URL."""
    try:
        max_bytes = 10_000_000  # 10MB limit
        with requests.get(
            url,
            timeout=30,
            headers={"User-Agent": "Agent-Gantry/0.1.0"},
            stream=True,
        ) as response:
            response.raise_for_status()

            # Validate content length header before downloading body
            content_length = response.headers.get("content-length")
            if content_length is not None:
                try:
                    if int(content_length) > max_bytes:
                        raise ValueError(f"Content too large: {content_length} bytes")
                except ValueError:
                    # Ignore invalid Content-Length value and fall back to streaming limit
                    pass

            # Stream response body and enforce maximum size while downloading
            chunks: list[bytes] = []
            bytes_read = 0
            for chunk in response.iter_content(chunk_size=8192):
                if not chunk:
                    continue
                chunks.append(chunk)
                bytes_read += len(chunk)
                if bytes_read > max_bytes:
                    raise ValueError(f"Content too large: {bytes_read} bytes")

            encoding = response.encoding or "utf-8"
            return b"".join(chunks).decode(encoding, errors="replace")
    except requests.exceptions.Timeout:
        raise ValueError(f"Request timed out while fetching: {url}")
    except requests.exceptions.ConnectionError:
        raise ValueError(f"Connection error while fetching: {url}")
    except requests.exceptions.HTTPError as e:
        raise ValueError(f"HTTP error while fetching {url}: {e}")
    except requests.exceptions.RequestException as e:
        raise ValueError(f"Error fetching {url}: {e}")

@tools.register(tags=["datetime"])
def get_current_utc_time() -> str:
    """Get the current UTC time as an ISO formatted string."""
    return datetime.datetime.now(datetime.timezone.utc).isoformat().replace("+00:00", "Z")

# --- File System Tools ---

@tools.register(tags=["fs", "file"])
def list_directory(path: str = ".") -> list[str]:
    """List the contents of a directory."""
    return os.listdir(path)

@tools.register(tags=["fs", "file"])
def read_text_file(path: str) -> str:
    """Read the contents of a text file."""
    with open(path, encoding='utf-8') as f:
        return f.read()

@tools.register(tags=["fs", "file"])
def write_text_file(path: str, content: str) -> str:
    """Write content to a text file."""
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)
    return f"File written to {path}"

@tools.register(tags=["fs", "file"])
def append_text_file(path: str, content: str) -> str:
    """Append content to a text file."""
    with open(path, 'a', encoding='utf-8') as f:
        f.write(content)
    return f"Content appended to {path}"

@tools.register(tags=["fs", "file"])
def delete_file(path: str) -> str:
    """Delete a file."""
    os.remove(path)
    return f"File {path} deleted"

@tools.register(tags=["fs", "directory"])
def create_directory(path: str) -> str:
    """Create a new directory."""
    os.makedirs(path, exist_ok=True)
    return f"Directory {path} created"

@tools.register(tags=["fs", "directory"])
def delete_directory(path: str) -> str:
    """Delete a directory and all its contents."""
    shutil.rmtree(path)
    return f"Directory {path} deleted"

@tools.register(tags=["fs", "file"])
def file_exists(path: str) -> bool:
    """Check if a file or directory exists."""
    return os.path.exists(path)

@tools.register(tags=["fs", "file"])
def get_file_size(path: str) -> int:
    """Get the size of a file in bytes."""
    return os.path.getsize(path)

@tools.register(tags=["fs", "file"])
def get_file_extension(path: str) -> str:
    """Get the extension of a file."""
    return pathlib.Path(path).suffix

@tools.register(tags=["fs", "file"])
def move_file(src: str, dst: str) -> str:
    """Move a file or directory to a new location."""
    shutil.move(src, dst)
    return f"Moved {src} to {dst}"

@tools.register(tags=["fs", "file"])
def copy_file(src: str, dst: str) -> str:
    """Copy a file to a new location."""
    shutil.copy2(src, dst)
    return f"Copied {src} to {dst}"

@tools.register(tags=["fs", "search"])
def search_files(pattern: str, root_dir: str = ".") -> list[str]:
    """Search for files matching a glob pattern."""
    return [str(p) for p in pathlib.Path(root_dir).rglob(pattern)]

@tools.register(tags=["fs", "temp"])
def get_temp_directory() -> str:
    """Get the path to the system temporary directory."""
    return tempfile.gettempdir()

# --- System & OS Tools ---

@tools.register(tags=["system", "os"])
def get_os_name() -> str:
    """Get the name of the operating system."""
    return os.name

@tools.register(tags=["system", "os"])
def get_platform_info() -> str:
    """Get detailed platform information."""
    return platform.platform()

@tools.register(tags=["system", "cpu"])
def get_cpu_count() -> int:
    """Get the number of logical CPUs in the system."""
    return os.cpu_count() or 0

@tools.register(tags=["system", "env"])
def get_environment_variable(name: str, default: str | None = None) -> str | None:
    """Get the value of an environment variable."""
    return os.environ.get(name, default)

@tools.register(tags=["system", "env"])
def set_environment_variable(name: str, value: str) -> str:
    """Set the value of an environment variable."""
    os.environ[name] = value
    return f"Environment variable {name} set"

@tools.register(tags=["system", "os"])
def get_current_working_directory() -> str:
    """Get the current working directory."""
    return os.getcwd()

@tools.register(tags=["system", "user"])
def get_user_name() -> str:
    """Get the name of the current user."""
    return os.getlogin()

@tools.register(tags=["system", "network"])
def get_hostname() -> str:
    """Get the hostname of the machine."""
    return socket.gethostname()

# --- Network Tools ---

@tools.register(tags=["network", "dns"])
def resolve_dns(hostname: str) -> str:
    """Resolve a hostname to an IP address."""
    return socket.gethostbyname(hostname)

@tools.register(tags=["network", "ip"])
def get_local_ip() -> str:
    """Get the local IP address of the machine."""
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        # doesn't even have to be reachable
        s.connect(('10.255.255.255', 1))
        ip = s.getsockname()[0]
    except Exception:
        ip = '127.0.0.1'
    finally:
        s.close()
    return ip

@tools.register(tags=["network", "port"])
def is_port_open(host: str, port: int, timeout: float = 1.0) -> bool:
    """Check if a specific port is open on a host."""
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.settimeout(timeout)
        return s.connect_ex((host, port)) == 0

@tools.register(tags=["network", "http"])
def http_get_status(url: str) -> int:
    """Get the HTTP status code of a URL."""
    response = requests.get(url, timeout=10)
    return response.status_code

@tools.register(tags=["network", "http"])
def http_post_json(url: str, data: dict[str, Any]) -> dict[str, Any]:
    """Send a POST request with JSON data."""
    response = requests.post(url, json=data, timeout=10)
    return response.json()

@tools.register(tags=["network", "url"])
def url_encode(text: str) -> str:
    """URL-encode a string."""
    import urllib.parse
    return urllib.parse.quote(text)

@tools.register(tags=["network", "url"])
def url_decode(text: str) -> str:
    """URL-decode a string."""
    import urllib.parse
    return urllib.parse.unquote(text)

# --- Data & Serialization Tools ---

@tools.register(tags=["data", "json"])
def json_to_dict(json_str: str) -> dict[str, Any]:
    """Convert a JSON string to a dictionary."""
    return json.loads(json_str)

@tools.register(tags=["data", "json"])
def dict_to_json(data: dict[str, Any], indent: int = 4) -> str:
    """Convert a dictionary to a JSON string."""
    return json.dumps(data, indent=indent)

@tools.register(tags=["data", "csv"])
def csv_to_list(csv_str: str) -> list[dict[str, str]]:
    """Convert a CSV string to a list of dictionaries."""
    import csv
    f = io.StringIO(csv_str)
    reader = csv.DictReader(f)
    return list(reader)

@tools.register(tags=["data", "base64"])
def base64_encode(text: str) -> str:
    """Encode a string to Base64."""
    return base64.b64encode(text.encode()).decode()

@tools.register(tags=["data", "base64"])
def base64_decode(encoded_str: str) -> str:
    """Decode a Base64 string."""
    return base64.b64decode(encoded_str.encode()).decode()

@tools.register(tags=["data", "uuid"])
def generate_uuid() -> str:
    """Generate a random UUID (v4)."""
    return str(uuid.uuid4())

@tools.register(tags=["data", "hash"])
def get_hash(text: str, algorithm: str = "sha256") -> str:
    """Calculate the hash of a string."""
    h = hashlib.new(algorithm)
    h.update(text.encode())
    return h.hexdigest()

@tools.register(tags=["data", "compression"])
def compress_string(text: str) -> str:
    """Compress a string using zlib and encode to Base64."""
    import zlib
    compressed = zlib.compress(text.encode())
    return base64.b64encode(compressed).decode()

@tools.register(tags=["data", "compression"])
def decompress_string(compressed_base64: str) -> str:
    """Decompress a Base64-encoded zlib string."""
    import zlib
    compressed = base64.b64decode(compressed_base64.encode())
    return zlib.decompress(compressed).decode()

# --- Text & NLP Tools ---

@tools.register(tags=["text", "regex"])
def regex_search(pattern: str, text: str) -> list[str]:
    """Search for all occurrences of a regex pattern in text."""
    return re.findall(pattern, text)

@tools.register(tags=["text", "regex"])
def regex_replace(pattern: str, replacement: str, text: str) -> str:
    """Replace occurrences of a regex pattern in text."""
    return re.sub(pattern, replacement, text)

@tools.register(tags=["text"])
def count_words(text: str) -> int:
    """Count the number of words in a string."""
    return len(text.split())

@tools.register(tags=["text"])
def count_characters(text: str, include_whitespace: bool = True) -> int:
    """Count the number of characters in a string."""
    if include_whitespace:
        return len(text)
    return len(text.replace(" ", "").replace("\n", "").replace("\t", ""))

@tools.register(tags=["text"])
def reverse_string(text: str) -> str:
    """Reverse a string."""
    return text[::-1]

@tools.register(tags=["text"])
def extract_emails(text: str) -> list[str]:
    """Extract all email addresses from text."""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    return re.findall(email_pattern, text)

@tools.register(tags=["text"])
def extract_urls(text: str) -> list[str]:
    """Extract all URLs from text."""
    url_pattern = r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+'
    return re.findall(url_pattern, text)

@tools.register(tags=["text"])
def strip_html_tags(html: str) -> str:
    """Remove HTML tags from a string."""
    return re.sub(r'<[^>]*>', '', html)

# --- Math & Statistics Tools ---

@tools.register(tags=["math", "stats"])
def calculate_mean(numbers: list[float]) -> float:
    """Calculate the arithmetic mean of a list of numbers."""
    return statistics.mean(numbers)

@tools.register(tags=["math", "stats"])
def calculate_median(numbers: list[float]) -> float:
    """Calculate the median of a list of numbers."""
    return statistics.median(numbers)

@tools.register(tags=["math", "stats"])
def calculate_stdev(numbers: list[float]) -> float:
    """Calculate the standard deviation of a list of numbers."""
    return statistics.stdev(numbers)

@tools.register(tags=["math"])
def calculate_factorial(n: int) -> int:
    """Calculate the factorial of a non-negative integer."""
    return math.factorial(n)

@tools.register(tags=["math"])
def is_prime(n: int) -> bool:
    """Check if a number is prime."""
    if n < 2: return False
    for i in range(2, int(math.sqrt(n)) + 1):
        if n % i == 0: return False
    return True

@tools.register(tags=["math", "random"])
def get_random_int(min_val: int, max_val: int) -> int:
    """Generate a random integer between min_val and max_val (inclusive)."""
    return random.randint(min_val, max_val)

@tools.register(tags=["math"])
def calculate_percentage(part: float, whole: float) -> float:
    """Calculate what percentage part is of whole."""
    if whole == 0: return 0.0
    return (part / whole) * 100

# --- Time & Productivity Tools ---

@tools.register(tags=["time"])
def get_current_timestamp() -> float:
    """Get the current Unix timestamp."""
    return time.time()

@tools.register(tags=["time"])
def format_timestamp(timestamp: float, format_str: str = "%Y-%m-%d %H:%M:%S") -> str:
    """Format a Unix timestamp into a human-readable string."""
    return datetime.datetime.fromtimestamp(timestamp).strftime(format_str)

@tools.register(tags=["time"])
def get_days_until(target_date: str) -> int:
    """Calculate the number of days from today until the target date (YYYY-MM-DD)."""
    today = datetime.date.today()
    target = datetime.datetime.strptime(target_date, "%Y-%m-%d").date()
    return (target - today).days

@tools.register(tags=["time"])
def is_leap_year(year: int) -> bool:
    """Check if a year is a leap year."""
    import calendar
    return calendar.isleap(year)

@tools.register(tags=["time"])
def get_weekday(date_str: str) -> str:
    """Get the day of the week for a given date (YYYY-MM-DD)."""
    d = datetime.datetime.strptime(date_str, "%Y-%m-%d")
    return d.strftime("%A")

@tools.register(tags=["productivity"])
def sleep(seconds: float) -> str:
    """Pause execution for a specified number of seconds."""
    time.sleep(seconds)
    return f"Slept for {seconds} seconds"

# --- Web & API Tools ---

@tools.register(tags=["web", "http"])
def get_http_headers(url: str) -> dict[str, str]:
    """Get the HTTP headers of a URL."""
    response = requests.head(url, timeout=10)
    return dict(response.headers)

@tools.register(tags=["web"])
def get_domain_from_url(url: str) -> str:
    """Extract the domain name from a URL."""
    from urllib.parse import urlparse
    return urlparse(url).netloc

@tools.register(tags=["web"])
def is_valid_url(url: str) -> bool:
    """Check if a string is a valid URL."""
    from urllib.parse import urlparse
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except ValueError:
        return False

# --- Security Tools ---

@tools.register(tags=["security"])
def generate_password(length: int = 12, include_special: bool = True) -> str:
    """Generate a random secure password."""
    import secrets
    import string
    alphabet = string.ascii_letters + string.digits
    if include_special:
        alphabet += string.punctuation
    return ''.join(secrets.choice(alphabet) for _ in range(length))

@tools.register(tags=["security"])
def is_strong_password(password: str) -> bool:
    """Check if a password meets basic strength requirements."""
    if len(password) < 8: return False
    if not any(c.isupper() for c in password): return False
    if not any(c.islower() for c in password): return False
    if not any(c.isdigit() for c in password): return False
    return True

@tools.register(tags=["security"])
def mask_sensitive_data(text: str, visible_chars: int = 4) -> str:
    """Mask sensitive data, leaving only the last few characters visible."""
    if len(text) <= visible_chars:
        return "*" * len(text)
    return "*" * (len(text) - visible_chars) + text[-visible_chars:]

# --- Misc Tools ---

@tools.register(tags=["misc"])
def get_random_color() -> str:
    """Generate a random hex color code."""
    return f"#{random.randint(0, 0xFFFFFF):06x}"

@tools.register(tags=["misc", "math"])
def convert_temperature(value: float, from_unit: str, to_unit: str) -> float:
    """Convert temperature between Celsius, Fahrenheit, and Kelvin."""
    from_unit = from_unit.upper()[0]
    to_unit = to_unit.upper()[0]

    # Convert to Celsius first
    if from_unit == 'F':
        c = (value - 32) * 5/9
    elif from_unit == 'K':
        c = value - 273.15
    else:
        c = value

    # Convert from Celsius to target
    if to_unit == 'F':
        return (c * 9/5) + 32
    elif to_unit == 'K':
        return c + 273.15
    else:
        return c

@tools.register(tags=["misc", "health"])
def calculate_bmi(weight_kg: float, height_m: float) -> float:
    """Calculate Body Mass Index (BMI)."""
    if height_m == 0: return 0.0
    return weight_kg / (height_m ** 2)

@tools.register(tags=["misc", "system"])
def get_python_version() -> str:
    """Get the current Python version."""
    return platform.python_version()

@tools.register(tags=["misc", "system"])
def get_installed_packages() -> list[str]:
    """Get a list of installed Python packages (names only)."""
    try:
        result = subprocess.run([sys.executable, "-m", "pip", "list", "--format=json"],
                               capture_output=True, text=True)
        if result.returncode == 0:
            packages = json.loads(result.stdout)
            return [p['name'] for p in packages]
    except Exception:
        pass
    return []

# --- Specialized Chemistry Tools ---

@tools.register(tags=["chemistry"])
def get_smiles_from_name(name: str) -> str:
    """Get the SMILES string for a compound name using PubChem."""
    compounds = pcp.get_compounds(name, 'name')
    if not compounds:
        raise ValueError(f"No compound found for name: {name}")
    return compounds[0].isomeric_smiles

@tools.register(tags=["chemistry"])
def calculate_logp(smiles: str) -> float:
    """Calculate the Octanol-Water Partition Coefficient (LogP) from SMILES."""
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        raise ValueError("Invalid SMILES string.")
    return Descriptors.MolLogP(mol)

@tools.register(tags=["chemistry"])
def calculate_tpsa(smiles: str) -> float:
    """Calculate the Topological Polar Surface Area (TPSA) from SMILES."""
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        raise ValueError("Invalid SMILES string.")
    return Descriptors.TPSA(mol)

@tools.register(tags=["chemistry"])
def is_valid_smiles(smiles: str) -> bool:
    """Check if a SMILES string is valid."""
    return Chem.MolFromSmiles(smiles) is not None

# --- Advanced Math & Geometry Tools ---

@tools.register(tags=["math", "geometry"])
def calculate_triangle_area(base: float, height: float) -> float:
    """Calculate the area of a triangle."""
    return 0.5 * base * height

@tools.register(tags=["math", "geometry"])
def calculate_sphere_volume(radius: float) -> float:
    """Calculate the volume of a sphere."""
    return (4/3) * math.pi * (radius ** 3)

@tools.register(tags=["math", "geometry"])
def calculate_distance_2d(x1: float, y1: float, x2: float, y2: float) -> float:
    """Calculate the Euclidean distance between two points in 2D space."""
    return math.sqrt((x2 - x1)**2 + (y2 - y1)**2)

@tools.register(tags=["math"])
def get_gcd(a: int, b: int) -> int:
    """Calculate the Greatest Common Divisor of two numbers."""
    return math.gcd(a, b)

@tools.register(tags=["math"])
def get_lcm(a: int, b: int) -> int:
    """Calculate the Least Common Multiple of two numbers."""
    if a == 0 or b == 0: return 0
    return abs(a * b) // math.gcd(a, b)

# --- Advanced Text Tools ---

@tools.register(tags=["text"])
def is_palindrome(text: str) -> bool:
    """Check if a string is a palindrome."""
    clean_text = re.sub(r'[^a-zA-Z0-9]', '', text).lower()
    return clean_text == clean_text[::-1]

@tools.register(tags=["text"])
def count_vowels(text: str) -> int:
    """Count the number of vowels in a string."""
    return sum(1 for char in text.lower() if char in 'aeiou')

@tools.register(tags=["text"])
def truncate_text(text: str, max_length: int, suffix: str = "...") -> str:
    """Truncate text to a maximum length with a suffix."""
    if len(text) <= max_length:
        return text
    return text[:max_length].rsplit(' ', 1)[0] + suffix

@tools.register(tags=["text"])
def camel_to_snake(text: str) -> str:
    """Convert CamelCase to snake_case."""
    s1 = re.sub('(.)([A-Z][a-z]+)', r'\1_\2', text)
    return re.sub('([a-z0-9])([A-Z])', r'\1_\2', s1).lower()

@tools.register(tags=["text"])
def snake_to_camel(text: str) -> str:
    """Convert snake_case to CamelCase."""
    return ''.join(word.title() for word in text.split('_'))

# --- Advanced System Tools ---

@tools.register(tags=["system"])
def run_shell_command(command: str) -> dict[str, Any]:
    """Run a shell command and return its output and exit code."""
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True, timeout=30)
        return {
            "stdout": result.stdout,
            "stderr": result.stderr,
            "exit_code": result.returncode
        }
    except Exception as e:
        return {"error": str(e)}

@tools.register(tags=["system"])
def get_process_id() -> int:
    """Get the current process ID."""
    return os.getpid()

@tools.register(tags=["system"])
def get_disk_usage(path: str = ".") -> dict[str, int]:
    """Get disk usage statistics for a path."""
    usage = shutil.disk_usage(path)
    return {
        "total": usage.total,
        "used": usage.used,
        "free": usage.free
    }

# --- Advanced Data Tools ---

@tools.register(tags=["data", "yaml"])
def yaml_to_dict(yaml_str: str) -> Any:
    """Convert a YAML string to a Python object."""
    import yaml
    return yaml.safe_load(yaml_str)

@tools.register(tags=["data", "yaml"])
def dict_to_yaml(data: Any) -> str:
    """Convert a Python object to a YAML string."""
    import yaml
    return yaml.dump(data, default_flow_style=False)

# --- Financial Tools ---

@tools.register(tags=["finance"])
def calculate_compound_interest(principal: float, rate: float, time: float, n: int = 1) -> float:
    """Calculate compound interest: A = P(1 + r/n)^(nt)."""
    return principal * (1 + rate/n)**(n*time)

@tools.register(tags=["finance"])
def calculate_loan_payment(principal: float, annual_rate: float, years: int) -> float:
    """Calculate monthly loan payment."""
    monthly_rate = annual_rate / 12 / 100
    n_payments = years * 12
    if monthly_rate == 0: return principal / n_payments
    return (principal * monthly_rate) / (1 - (1 + monthly_rate)**-n_payments)

@tools.register(tags=["finance"])
def calculate_roi(gain: float, cost: float) -> float:
    """Calculate Return on Investment (ROI) percentage."""
    if cost == 0: return 0.0
    return ((gain - cost) / cost) * 100

# --- Specific Unit Conversion Tools ---

@tools.register(tags=["conversion"])
def meters_to_feet(meters: float) -> float:
    """Convert meters to feet."""
    return meters * 3.28084

@tools.register(tags=["conversion"])
def feet_to_meters(feet: float) -> float:
    """Convert feet to meters."""
    return feet / 3.28084

@tools.register(tags=["conversion"])
def kilograms_to_pounds(kg: float) -> float:
    """Convert kilograms to pounds."""
    return kg * 2.20462

@tools.register(tags=["conversion"])
def pounds_to_kilograms(lbs: float) -> float:
    """Convert pounds to kilograms."""
    return lbs / 2.20462

@tools.register(tags=["conversion"])
def km_to_miles(km: float) -> float:
    """Convert kilometers to miles."""
    return km * 0.621371

@tools.register(tags=["conversion"])
def miles_to_km(miles: float) -> float:
    """Convert miles to kilometers."""
    return miles / 0.621371

# --- More Math Tools ---

@tools.register(tags=["math"])
def calculate_logarithm(x: float, base: float = math.e) -> float:
    """Calculate the logarithm of x to the given base."""
    return math.log(x, base)

@tools.register(tags=["math"])
def calculate_power(base: float, exponent: float) -> float:
    """Calculate base raised to the power of exponent."""
    return math.pow(base, exponent)

@tools.register(tags=["math"])
def calculate_sine(angle_degrees: float) -> float:
    """Calculate the sine of an angle in degrees."""
    return math.sin(math.radians(angle_degrees))

@tools.register(tags=["math"])
def calculate_cosine(angle_degrees: float) -> float:
    """Calculate the cosine of an angle in degrees."""
    return math.cos(math.radians(angle_degrees))

@tools.register(tags=["math"])
def calculate_tangent(angle_degrees: float) -> float:
    """Calculate the tangent of an angle in degrees."""
    return math.tan(math.radians(angle_degrees))

# --- More Text Tools ---

@tools.register(tags=["text"])
def count_sentences(text: str) -> int:
    """Count the number of sentences in a string."""
    sentences = re.split(r'[.!?]+', text)
    return len([s for s in sentences if s.strip()])

@tools.register(tags=["text"])
def remove_punctuation(text: str) -> str:
    """Remove all punctuation from a string."""
    import string
    return text.translate(str.maketrans('', '', string.punctuation))

@tools.register(tags=["text"])
def get_word_frequency(text: str) -> dict[str, int]:
    """Get the frequency of each word in a string."""
    words = remove_punctuation(text).lower().split()
    freq = {}
    for word in words:
        freq[word] = freq.get(word, 0) + 1
    return freq

@tools.register(tags=["text"])
def is_numeric(text: str) -> bool:
    """Check if a string contains only numeric characters."""
    return text.isnumeric()

@tools.register(tags=["text"])
def is_alphabetic(text: str) -> bool:
    """Check if a string contains only alphabetic characters."""
    return text.isalpha()

# --- More File System Tools ---

@tools.register(tags=["fs"])
def get_file_modification_time(path: str) -> str:
    """Get the last modification time of a file."""
    mtime = os.path.getmtime(path)
    return datetime.datetime.fromtimestamp(mtime).isoformat()

@tools.register(tags=["fs"])
def is_file(path: str) -> bool:
    """Check if a path is a file."""
    return os.path.isfile(path)

@tools.register(tags=["fs"])
def is_directory(path: str) -> bool:
    """Check if a path is a directory."""
    return os.path.isdir(path)

@tools.register(tags=["fs"])
def get_absolute_path(path: str) -> str:
    """Get the absolute path of a file or directory."""
    return os.path.abspath(path)

@tools.register(tags=["fs"])
def join_paths(*parts: str) -> str:
    """Join multiple path components into one."""
    return os.path.join(*parts)

# --- More Network Tools ---

@tools.register(tags=["network"])
def is_ipv4(ip: str) -> bool:
    """Check if a string is a valid IPv4 address."""
    try:
        socket.inet_aton(ip)
        return True
    except OSError:
        return False

@tools.register(tags=["network"])
def get_ip_from_hostname(hostname: str) -> str:
    """Get the IP address associated with a hostname."""
    return socket.gethostbyname(hostname)

# --- Advanced Data Manipulation Tools ---

@tools.register(tags=["data"])
def flatten_dict(d: dict[str, Any], parent_key: str = '', sep: str = '.') -> dict[str, Any]:
    """Flatten a nested dictionary."""
    items: list[tuple[str, Any]] = []
    for k, v in d.items():
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.extend(flatten_dict(v, new_key, sep=sep).items())
        else:
            items.append((new_key, v))
    return dict(items)

@tools.register(tags=["data"])
def merge_dicts(dict1: dict[str, Any], dict2: dict[str, Any]) -> dict[str, Any]:
    """Merge two dictionaries."""
    res = dict1.copy()
    res.update(dict2)
    return res

@tools.register(tags=["data"])
def sort_dict_by_key(d: dict[str, Any]) -> dict[str, Any]:
    """Sort a dictionary by its keys."""
    return dict(sorted(d.items()))

# --- Encoding & Decoding Tools ---

@tools.register(tags=["encoding"])
def to_binary(text: str) -> str:
    """Convert a string to its binary representation."""
    return ' '.join(format(ord(c), '08b') for c in text)

@tools.register(tags=["encoding"])
def from_binary(binary_str: str) -> str:
    """Convert a binary string back to text."""
    binary_values = binary_str.split()
    return ''.join(chr(int(bv, 2)) for bv in binary_values)

@tools.register(tags=["encoding"])
def to_hex(text: str) -> str:
    """Convert a string to its hexadecimal representation."""
    return text.encode().hex()

@tools.register(tags=["encoding"])
def from_hex(hex_str: str) -> str:
    """Convert a hexadecimal string back to text."""
    return bytes.fromhex(hex_str).decode()

@tools.register(tags=["encoding"])
def rot13_encode(text: str) -> str:
    """Encode a string using ROT13."""
    import codecs
    return codecs.encode(text, 'rot_13')

# --- Advanced Math & Probability Tools ---

@tools.register(tags=["math"])
def calculate_combinations(n: int, k: int) -> int:
    """Calculate the number of combinations (n choose k)."""
    return math.comb(n, k)

@tools.register(tags=["math"])
def calculate_permutations(n: int, k: int) -> int:
    """Calculate the number of permutations (nPk)."""
    return math.perm(n, k)

@tools.register(tags=["math"])
def is_even(n: int) -> bool:
    """Check if a number is even."""
    return n % 2 == 0

@tools.register(tags=["math"])
def is_odd(n: int) -> bool:
    """Check if a number is odd."""
    return n % 2 != 0

@tools.register(tags=["math"])
def get_divisors(n: int) -> list[int]:
    """Get all divisors of a number."""
    divisors = []
    for i in range(1, int(math.sqrt(n)) + 1):
        if n % i == 0:
            divisors.append(i)
            if i*i != n:
                divisors.append(n // i)
    return sorted(divisors)

# --- System Information Tools ---

@tools.register(tags=["system"])
def get_system_architecture() -> str:
    """Get the system architecture (e.g., 'x86_64')."""
    return platform.machine()

@tools.register(tags=["system"])
def get_system_version() -> str:
    """Get the system version string."""
    return platform.version()

# --- Detailed Time Tools ---

@tools.register(tags=["time"])
def get_current_year() -> int:
    """Get the current year."""
    return datetime.datetime.now().year

@tools.register(tags=["time"])
def get_current_month() -> int:
    """Get the current month (1-12)."""
    return datetime.datetime.now().month

@tools.register(tags=["time"])
def get_current_day() -> int:
    """Get the current day of the month (1-31)."""
    return datetime.datetime.now().day

@tools.register(tags=["time"])
def get_time_difference_seconds(start_time: float, end_time: float) -> float:
    """Calculate the difference between two timestamps in seconds."""
    return end_time - start_time

# --- URL Manipulation Tools ---

@tools.register(tags=["web", "url"])
def get_url_query_params(url: str) -> dict[str, list[str]]:
    """Extract query parameters from a URL."""
    from urllib.parse import parse_qs, urlparse
    return parse_qs(urlparse(url).query)

@tools.register(tags=["web", "url"])
def get_url_path(url: str) -> str:
    """Extract the path component from a URL."""
    from urllib.parse import urlparse
    return urlparse(url).path

# --- List Manipulation Tools ---

@tools.register(tags=["misc", "list"])
def get_random_element(elements: list[Any]) -> Any:
    """Get a random element from a list."""
    if not elements: return None
    return random.choice(elements)

@tools.register(tags=["misc", "list"])
def shuffle_list(elements: list[Any]) -> list[Any]:
    """Shuffle a list in place and return it."""
    res = elements.copy()
    random.shuffle(res)
    return res

@tools.register(tags=["misc", "list"])
def get_unique_elements(elements: list[Any]) -> list[Any]:
    """Get unique elements from a list, preserving order."""
    seen = set()
    return [x for x in elements if not (x in seen or seen.add(x))]

@tools.register(tags=["misc", "list"])
def count_occurrences(elements: list[Any], value: Any) -> int:
    """Count the number of occurrences of a value in a list."""
    return elements.count(value)

# --- Advanced Text Formatting Tools ---

@tools.register(tags=["text"])
def to_kebab_case(text: str) -> str:
    """Convert a string to kebab-case."""
    return re.sub(r'[\s_]+', '-', text).lower()

@tools.register(tags=["text"])
def remove_extra_whitespace(text: str) -> str:
    """Remove extra whitespace from a string."""
    return ' '.join(text.split())

@tools.register(tags=["text"])
def get_longest_word(text: str) -> str:
    """Get the longest word in a string."""
    words = remove_punctuation(text).split()
    if not words: return ""
    return max(words, key=len)

# --- Advanced Geometry & Number Theory Tools ---

@tools.register(tags=["math", "geometry"])
def calculate_cone_volume(radius: float, height: float) -> float:
    """Calculate the volume of a cone."""
    return (1/3) * math.pi * (radius ** 2) * height

@tools.register(tags=["math", "geometry"])
def calculate_ellipse_area(a: float, b: float) -> float:
    """Calculate the area of an ellipse."""
    return math.pi * a * b

@tools.register(tags=["math"])
def is_perfect_square(n: int) -> bool:
    """Check if a number is a perfect square."""
    if n < 0: return False
    sqrt_n = int(math.isqrt(n))
    return sqrt_n * sqrt_n == n

@tools.register(tags=["math"])
def get_fibonacci_number(n: int) -> int:
    """Get the n-th Fibonacci number."""
    if n <= 0: return 0
    if n == 1: return 1
    a, b = 0, 1
    for _ in range(2, n + 1):
        a, b = b, a + b
    return b

# --- Data Validation Tools ---

@tools.register(tags=["data", "validation"])
def is_valid_json(json_str: str) -> bool:
    """Check if a string is valid JSON."""
    try:
        json.loads(json_str)
        return True
    except ValueError:
        return False

@tools.register(tags=["data"])
def get_dict_depth(d: Any) -> int:
    """Get the maximum depth of a nested dictionary."""
    if not isinstance(d, dict) or not d:
        return 0
    return 1 + max(get_dict_depth(v) for v in d.values())

# --- System & Platform Tools ---

@tools.register(tags=["system"])
def get_system_platform() -> str:
    """Get the system platform (e.g., 'linux', 'win32', 'darwin')."""
    return sys.platform

@tools.register(tags=["system"])
def get_system_processor() -> str:
    """Get the system processor name."""
    return platform.processor()

# --- Advanced Time & Calendar Tools ---

@tools.register(tags=["time"])
def get_current_week_number() -> int:
    """Get the current ISO week number."""
    return datetime.datetime.now().isocalendar()[1]

@tools.register(tags=["time"])
def get_current_day_of_year() -> int:
    """Get the current day of the year (1-366)."""
    return datetime.datetime.now().timetuple().tm_yday

@tools.register(tags=["time"])
def is_weekend(date_str: str | None = None) -> bool:
    """Check if a date is a weekend. If no date is provided, checks today."""
    if date_str:
        d = datetime.datetime.strptime(date_str, "%Y-%m-%d")
    else:
        d = datetime.datetime.now()
    return d.weekday() >= 5

# --- Advanced URL Tools ---

@tools.register(tags=["web", "url"])
def is_secure_url(url: str) -> bool:
    """Check if a URL uses HTTPS."""
    from urllib.parse import urlparse
    return urlparse(url).scheme == 'https'

@tools.register(tags=["web", "url"])
def get_url_port(url: str) -> int | None:
    """Extract the port number from a URL."""
    from urllib.parse import urlparse
    return urlparse(url).port

# --- Random Generation Tools ---

@tools.register(tags=["misc", "random"])
def get_random_boolean() -> bool:
    """Generate a random boolean value."""
    return random.choice([True, False])

@tools.register(tags=["misc", "random"])
def get_random_ip() -> str:
    """Generate a random IPv4 address."""
    return ".".join(map(str, (random.randint(0, 255) for _ in range(4))))

@tools.register(tags=["misc", "random"])
def get_random_mac_address() -> str:
    """Generate a random MAC address."""
    return ":".join(f"{random.randint(0, 255):02x}" for _ in range(6))

# --- Advanced Line-Based Text Tools ---

@tools.register(tags=["text"])
def count_lines(text: str) -> int:
    """Count the number of lines in a string."""
    return len(text.splitlines())

@tools.register(tags=["text"])
def get_first_line(text: str) -> str:
    """Get the first line of a string."""
    lines = text.splitlines()
    return lines[0] if lines else ""

@tools.register(tags=["text"])
def sort_lines(text: str, reverse: bool = False) -> str:
    """Sort the lines of a string."""
    lines = text.splitlines()
    return "\n".join(sorted(lines, reverse=reverse))

# --- Advanced Statistical Tools ---

@tools.register(tags=["math", "stats"])
def calculate_mean_absolute_error(actual: list[float], predicted: list[float]) -> float:
    """Calculate the Mean Absolute Error (MAE)."""
    if len(actual) != len(predicted) or not actual:
        raise ValueError("Lists must be of the same non-zero length.")
    return sum(abs(a - p) for a, p in zip(actual, predicted)) / len(actual)

@tools.register(tags=["math", "stats"])
def calculate_z_score(value: float, mean: float, stdev: float) -> float:
    """Calculate the Z-score of a value."""
    if stdev == 0: return 0.0
    return (value - mean) / stdev

# --- Advanced Dictionary Path Tools ---

@tools.register(tags=["data"])
def get_dict_value_by_path(d: dict[str, Any], path: str, sep: str = '.') -> Any:
    """Get a value from a nested dictionary using a path string."""
    keys = path.split(sep)
    val = d
    for key in keys:
        if isinstance(val, dict) and key in val:
            val = val[key]
        else:
            return None
    return val

@tools.register(tags=["data"])
def has_dict_path(d: dict[str, Any], path: str, sep: str = '.') -> bool:
    """Check if a path exists in a nested dictionary."""
    keys = path.split(sep)
    val = d
    for key in keys:
        if isinstance(val, dict) and key in val:
            val = val[key]
        else:
            return False
    return True

# --- Mock System Resource Tools ---

@tools.register(tags=["system"])
def get_mock_system_load() -> float:
    """Get a mock system load average (0.0 to 1.0)."""
    return round(random.uniform(0.1, 0.9), 2)

@tools.register(tags=["system"])
def get_mock_memory_usage() -> dict[str, float]:
    """Get mock memory usage statistics in GB."""
    total = 16.0
    used = round(random.uniform(4.0, 12.0), 1)
    return {"total": total, "used": used, "free": total - used}

# --- Advanced Time Difference Tools ---

@tools.register(tags=["time"])
def get_time_until_next_hour() -> float:
    """Get the number of seconds until the start of the next hour."""
    now = datetime.datetime.now()
    next_hour = (now + datetime.timedelta(hours=1)).replace(minute=0, second=0, microsecond=0)
    return (next_hour - now).total_seconds()

@tools.register(tags=["time"])
def get_time_since_last_day() -> float:
    """Get the number of seconds since the start of the current day."""
    now = datetime.datetime.now()
    start_of_day = now.replace(hour=0, minute=0, second=0, microsecond=0)
    return (now - start_of_day).total_seconds()

# --- Advanced URL Component Tools ---

@tools.register(tags=["web", "url"])
def get_url_extension(url: str) -> str:
    """Extract the file extension from a URL path."""
    from urllib.parse import urlparse
    path = urlparse(url).path
    return pathlib.Path(path).suffix

@tools.register(tags=["web", "url"])
def is_localhost_url(url: str) -> bool:
    """Check if a URL points to localhost."""
    from urllib.parse import urlparse
    netloc = urlparse(url).netloc.split(':')[0]
    return netloc in ['localhost', '127.0.0.1', '::1']

# --- Mock Data Generation Tools ---

@tools.register(tags=["misc", "random"])
def get_random_phone_number() -> str:
    """Generate a random mock phone number."""
    return f"+1-{random.randint(200, 999)}-{random.randint(200, 999)}-{random.randint(1000, 9999)}"

@tools.register(tags=["misc", "random"])
def get_random_email_address() -> str:
    """Generate a random mock email address."""
    domains = ["example.com", "test.org", "demo.net", "mail.io"]
    name = "".join(random.choices("abcdefghijklmnopqrstuvwxyz", k=8))
    return f"{name}@{random.choice(domains)}"

@tools.register(tags=["misc", "random"])
def get_random_sentence() -> str:
    """Generate a random mock sentence."""
    subjects = ["The cat", "A dog", "The programmer", "An agent", "The system"]
    verbs = ["runs", "jumps", "codes", "thinks", "executes"]
    objects = ["quickly", "efficiently", "with joy", "the task", "the code"]
    return f"{random.choice(subjects)} {random.choice(verbs)} {random.choice(objects)}."

# --- Advanced Text Search & Analysis Tools ---

@tools.register(tags=["text"])
def get_all_indices(text: str, sub: str) -> list[int]:
    """Find all start indices of a substring in text."""
    indices = []
    start = 0
    while True:
        start = text.find(sub, start)
        if start == -1: return indices
        indices.append(start)
        start += 1

@tools.register(tags=["text"])
def is_ascii(text: str) -> bool:
    """Check if a string contains only ASCII characters."""
    return all(ord(c) < 128 for c in text)

# --- Advanced Statistical Mean Tools ---

@tools.register(tags=["math", "stats"])
def calculate_weighted_mean(values: list[float], weights: list[float]) -> float:
    """Calculate the weighted arithmetic mean."""
    if len(values) != len(weights) or not values:
        raise ValueError("Lists must be of the same non-zero length.")
    return sum(v * w for v, w in zip(values, weights)) / sum(weights)

@tools.register(tags=["math", "stats"])
def calculate_geometric_mean(numbers: list[float]) -> float:
    """Calculate the geometric mean of a list of numbers."""
    return statistics.geometric_mean(numbers)

# --- Advanced Dictionary Analysis Tools ---

@tools.register(tags=["data"])
def get_dict_key_with_max_value(d: dict[Any, float]) -> Any | None:
    """Get the key associated with the maximum value in a dictionary."""
    if not d: return None
    return max(d, key=d.get)

@tools.register(tags=["data"])
def invert_dict(d: dict[Any, Any]) -> dict[Any, list[Any]]:
    """Invert a dictionary, mapping values to lists of keys."""
    res: dict[Any, list[Any]] = {}
    for k, v in d.items():
        res.setdefault(v, []).append(k)
    return res

# --- Advanced System Configuration Tools ---

@tools.register(tags=["system"])
def get_system_endianness() -> str:
    """Get the system byte order ('little' or 'big')."""
    return sys.byteorder

@tools.register(tags=["system"])
def get_system_recursion_limit() -> int:
    """Get the current recursion limit."""
    return sys.getrecursionlimit()

# --- Advanced Calendar Tools ---

@tools.register(tags=["time"])
def get_days_in_month(year: int, month: int) -> int:
    """Get the number of days in a specific month."""
    import calendar
    return calendar.monthrange(year, month)[1]

@tools.register(tags=["time"])
def get_last_day_of_month(year: int, month: int) -> str:
    """Get the last day of a month as a string (YYYY-MM-DD)."""
    days = get_days_in_month(year, month)
    return f"{year:04d}-{month:02d}-{days:02d}"

# --- Advanced URL Parsing Tools ---

@tools.register(tags=["web", "url"])
def get_url_netloc(url: str) -> str:
    """Extract the network location (domain + port) from a URL."""
    from urllib.parse import urlparse
    return urlparse(url).netloc

@tools.register(tags=["web", "url"])
def get_url_scheme(url: str) -> str:
    """Extract the scheme (e.g., 'http', 'https') from a URL."""
    from urllib.parse import urlparse
    return urlparse(url).scheme

# --- Advanced Random Color Tools ---

@tools.register(tags=["misc", "random"])
def get_random_rgb_color() -> tuple[int, int, int]:
    """Generate a random RGB color tuple."""
    return (random.randint(0, 255), random.randint(0, 255), random.randint(0, 255))

@tools.register(tags=["misc", "random"])
def get_random_country_code() -> str:
    """Generate a random mock ISO country code."""
    codes = ["US", "GB", "CA", "DE", "FR", "JP", "CN", "BR", "AU", "IN"]
    return random.choice(codes)

@tools.register(tags=["misc", "random"])
def get_random_currency_code() -> str:
    """Generate a random mock ISO currency code."""
    codes = ["USD", "EUR", "GBP", "JPY", "CAD", "AUD", "CHF", "CNY", "INR"]
    return random.choice(codes)

# --- Data Science & Numerical Tools (Requires pandas, numpy) ---

@tools.register(tags=["data", "pandas"])
def create_dataframe_summary(data: list[dict[str, Any]]) -> dict[str, Any]:
    """Create a statistical summary of a list of dictionaries using Pandas."""
    import pandas as pd
    df = pd.DataFrame(data)
    return df.describe().to_dict()

@tools.register(tags=["data", "pandas"])
def filter_dataframe(data: list[dict[str, Any]], column: str, value: Any) -> list[dict[str, Any]]:
    """Filter a list of dictionaries based on a column value using Pandas."""
    import pandas as pd
    df = pd.DataFrame(data)
    filtered_df = df[df[column] == value]
    return filtered_df.to_dict(orient='records')

@tools.register(tags=["data", "numpy"])
def calculate_matrix_inverse(matrix: list[list[float]]) -> list[list[float]]:
    """Calculate the inverse of a square matrix using NumPy."""
    import numpy as np
    arr = np.array(matrix)
    inv = np.linalg.inv(arr)
    return inv.tolist()

@tools.register(tags=["data", "numpy"])
def calculate_eigenvalues(matrix: list[list[float]]) -> list[float]:
    """Calculate the eigenvalues of a square matrix using NumPy."""
    import numpy as np
    arr = np.array(matrix)
    eigenvalues, _ = np.linalg.eig(arr)
    return eigenvalues.tolist()

@tools.register(tags=["data", "numpy"])
def generate_normal_distribution(mean: float, std: float, size: int) -> list[float]:
    """Generate a list of numbers following a normal distribution."""
    import numpy as np
    return np.random.normal(mean, std, size).tolist()

# --- Machine Learning Tools (Requires scikit-learn) ---

@tools.register(tags=["ml", "sklearn"])
def train_simple_linear_regression(x: list[float], y: list[float]) -> dict[str, float]:
    """Train a simple linear regression model and return coefficients."""
    import numpy as np
    from sklearn.linear_model import LinearRegression
    X = np.array(x).reshape(-1, 1)
    Y = np.array(y)
    model = LinearRegression().fit(X, Y)
    return {"coefficient": float(model.coef_[0]), "intercept": float(model.intercept_)}

@tools.register(tags=["ml", "sklearn"])
def cluster_data_kmeans(data: list[list[float]], n_clusters: int = 3) -> list[int]:
    """Cluster data points using K-Means algorithm."""
    import numpy as np
    from sklearn.cluster import KMeans
    X = np.array(data)
    kmeans = KMeans(n_clusters=n_clusters, n_init='auto').fit(X)
    return kmeans.labels_.tolist()

# --- Image Processing Tools (Requires Pillow) ---

@tools.register(tags=["image", "pillow"])
def get_image_metadata(image_path: str) -> dict[str, Any]:
    """Get metadata (size, format, mode) of an image file."""
    from PIL import Image
    with Image.open(image_path) as img:
        return {
            "size": img.size,
            "format": img.format,
            "mode": img.mode,
            "info": img.info
        }

@tools.register(tags=["image", "pillow"])
def resize_image(image_path: str, output_path: str, width: int, height: int) -> str:
    """Resize an image and save it to a new path."""
    from PIL import Image
    with Image.open(image_path) as img:
        resized = img.resize((width, height))
        resized.save(output_path)
    return f"Image resized and saved to {output_path}"

@tools.register(tags=["image", "pillow"])
def convert_image_format(image_path: str, output_path: str, format: str) -> str:
    """Convert an image to a different format (e.g., PNG to JPEG)."""
    from PIL import Image
    with Image.open(image_path) as img:
        img.save(output_path, format=format)
    return f"Image converted and saved to {output_path}"

@tools.register(tags=["image", "pillow"])
def rotate_image(image_path: str, output_path: str, degrees: float) -> str:
    """Rotate an image by a specified number of degrees."""
    from PIL import Image
    with Image.open(image_path) as img:
        rotated = img.rotate(degrees, expand=True)
        rotated.save(output_path)
    return f"Image rotated and saved to {output_path}"

@tools.register(tags=["image", "pillow"])
def apply_grayscale_filter(image_path: str, output_path: str) -> str:
    """Convert an image to grayscale."""
    from PIL import Image
    with Image.open(image_path) as img:
        grayscale = img.convert("L")
        grayscale.save(output_path)
    return f"Grayscale image saved to {output_path}"

@tools.register(tags=["image", "pillow"])
def generate_thumbnail(image_path: str, output_path: str, size: tuple[int, int] = (128, 128)) -> str:
    """Generate a thumbnail for an image."""
    from PIL import Image
    with Image.open(image_path) as img:
        img.thumbnail(size)
        img.save(output_path)
    return f"Thumbnail saved to {output_path}"

# --- Document Processing Tools (Requires PyPDF2, python-docx) ---

@tools.register(tags=["document", "pdf"])
def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract all text from a PDF file."""
    import PyPDF2
    text = ""
    with open(pdf_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

@tools.register(tags=["document", "pdf"])
def get_pdf_page_count(pdf_path: str) -> int:
    """Get the number of pages in a PDF file."""
    import PyPDF2
    with open(pdf_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        return len(reader.pages)

@tools.register(tags=["document", "docx"])
def read_docx_file(docx_path: str) -> str:
    """Read the text content of a .docx file."""
    import docx
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

@tools.register(tags=["document", "docx"])
def create_docx_file(docx_path: str, content: str, title: str | None = None) -> str:
    """Create a new .docx file with the specified content."""
    import docx
    doc = docx.Document()
    if title:
        doc.add_heading(title, 0)
    doc.add_paragraph(content)
    doc.save(docx_path)
    return f"Document saved to {docx_path}"

# --- Natural Language Processing Tools (Requires textblob) ---

@tools.register(tags=["nlp", "textblob"])
def analyze_sentiment(text: str) -> dict[str, float]:
    """Analyze the sentiment of a text (polarity and subjectivity)."""
    from textblob import TextBlob
    blob = TextBlob(text)
    return {
        "polarity": blob.sentiment.polarity,
        "subjectivity": blob.sentiment.subjectivity
    }

@tools.register(tags=["nlp", "textblob"])
def translate_text(text: str, to_lang: str) -> str:
    """Translate text to a target language using TextBlob (requires internet)."""
    from textblob import TextBlob
    blob = TextBlob(text)
    return str(blob.translate(to=to_lang))

@tools.register(tags=["nlp", "textblob"])
def correct_spelling(text: str) -> str:
    """Correct the spelling of a text using TextBlob."""
    from textblob import TextBlob
    blob = TextBlob(text)
    return str(blob.correct())

# --- Visualization Tools (Requires matplotlib) ---

@tools.register(tags=["viz", "matplotlib"])
def generate_bar_chart_base64(labels: list[str], values: list[float], title: str = "Bar Chart") -> str:
    """Generate a bar chart and return it as a Base64-encoded PNG string."""
    import base64

    import matplotlib.pyplot as plt

    plt.figure(figsize=(10, 6))
    plt.bar(labels, values)
    plt.title(title)

    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close()
    buf.seek(0)
    return base64.b64encode(buf.read()).decode('utf-8')

@tools.register(tags=["viz", "matplotlib"])
def generate_line_plot_base64(x: list[float], y: list[float], title: str = "Line Plot") -> str:
    """Generate a line plot and return it as a Base64-encoded PNG string."""
    import base64

    import matplotlib.pyplot as plt

    plt.figure(figsize=(10, 6))
    plt.plot(x, y)
    plt.title(title)

    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close()
    buf.seek(0)
    return base64.b64encode(buf.read()).decode('utf-8')

@tools.register(tags=["viz", "matplotlib"])
def generate_pie_chart_base64(labels: list[str], values: list[float], title: str = "Pie Chart") -> str:
    """Generate a pie chart and return it as a Base64-encoded PNG string."""
    import base64

    import matplotlib.pyplot as plt

    plt.figure(figsize=(8, 8))
    plt.pie(values, labels=labels, autopct="%1.1f%%")
    plt.title(title)

    buf = io.BytesIO()
    plt.savefig(buf, format="png")
    plt.close()
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")

# --- Advanced System Monitoring Tools (Requires psutil) ---

@tools.register(tags=["system", "psutil"])
def get_cpu_usage_percent(interval: float = 1.0) -> float:
    """Get the current CPU usage percentage."""
    import psutil
    return psutil.cpu_percent(interval=interval)

@tools.register(tags=["system", "psutil"])
def get_memory_info() -> dict[str, Any]:
    """Get detailed system memory usage statistics."""
    import psutil
    mem = psutil.virtual_memory()
    return {
        "total": mem.total,
        "available": mem.available,
        "percent": mem.percent,
        "used": mem.used,
        "free": mem.free
    }

@tools.register(tags=["system", "psutil"])
def get_disk_partitions() -> list[dict[str, str]]:
    """Get a list of all disk partitions."""
    import psutil
    partitions = psutil.disk_partitions()
    return [{"device": p.device, "mountpoint": p.mountpoint, "fstype": p.fstype} for p in partitions]

@tools.register(tags=["system", "psutil"])
def get_running_processes(limit: int = 10) -> list[dict[str, Any]]:
    """Get a list of currently running processes."""
    import psutil
    processes = []
    for proc in psutil.process_iter(['pid', 'name', 'username']):
        try:
            processes.append(proc.info)
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            pass
        if len(processes) >= limit:
            break
    return processes

# --- Advanced Security & Cryptography Tools (Requires cryptography) ---

@tools.register(tags=["security", "cryptography"])
def generate_encryption_key() -> str:
    """Generate a new Fernet encryption key."""
    from cryptography.fernet import Fernet
    return Fernet.generate_key().decode()

@tools.register(tags=["security", "cryptography"])
def encrypt_message(message: str, key: str) -> str:
    """Encrypt a message using a Fernet key."""
    from cryptography.fernet import Fernet
    f = Fernet(key.encode())
    return f.encrypt(message.encode()).decode()

@tools.register(tags=["security", "cryptography"])
def decrypt_message(encrypted_message: str, key: str) -> str:
    """Decrypt a message using a Fernet key."""
    from cryptography.fernet import Fernet
    f = Fernet(key.encode())
    return f.decrypt(encrypted_message.encode()).decode()

# --- Web Scraping Tools (Requires beautifulsoup4) ---

@tools.register(tags=["web", "scraping"])
def extract_all_links_from_url(url: str) -> list[str]:
    """Extract all unique absolute links from a web page."""
    from urllib.parse import urljoin

    from bs4 import BeautifulSoup
    response = requests.get(url, timeout=10)
    soup = BeautifulSoup(response.text, 'html.parser')
    links = set()
    for a in soup.find_all('a', href=True):
        links.add(urljoin(url, a['href']))
    return list(links)

@tools.register(tags=["web", "scraping"])
def get_page_title(url: str) -> str:
    """Get the title of a web page."""
    from bs4 import BeautifulSoup
    response = requests.get(url, timeout=10)
    soup = BeautifulSoup(response.text, 'html.parser')
    return soup.title.string if soup.title else "No title found"

# --- Advanced Timezone Tools (Requires pytz) ---

@tools.register(tags=["time", "pytz"])
def get_time_in_timezone(timezone_name: str) -> str:
    """Get the current time in a specific timezone (e.g., 'America/New_York')."""
    from datetime import datetime

    import pytz
    tz = pytz.timezone(timezone_name)
    return datetime.now(tz).isoformat()

@tools.register(tags=["time", "pytz"])
def list_all_timezones() -> list[str]:
    """List all available timezone names."""
    import pytz
    return pytz.all_timezones

# --- Advanced Misc Tools ---

@tools.register(tags=["misc", "qrcode"])
def generate_qr_code_base64(data: str) -> str:
    """Generate a QR code for the given data and return it as a Base64-encoded PNG string."""
    import base64

    import qrcode

    img = qrcode.make(data)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return base64.b64encode(buf.read()).decode('utf-8')

@tools.register(tags=["misc", "barcode"])
def generate_barcode_base64(data: str) -> str:
    """Generate a barcode (Code128) for the given data and return it as a Base64-encoded PNG string."""
    import base64

    import barcode
    from barcode.writer import ImageWriter

    EAN = barcode.get_generator('code128')
    ean = EAN(data, writer=ImageWriter())
    buf = io.BytesIO()
    ean.write(buf)
    return base64.b64encode(buf.read()).decode('utf-8')

@tools.register(tags=["misc", "weather"])
def get_mock_weather(city: str) -> dict[str, Any]:
    """Get mock weather data for a city."""
    conditions = ["Sunny", "Cloudy", "Rainy", "Partly Cloudy", "Stormy"]
    return {
        "city": city,
        "temperature": round(random.uniform(-10, 40), 1),
        "condition": random.choice(conditions),
        "humidity": random.randint(10, 90),
        "wind_speed": round(random.uniform(0, 100), 1)
    }

@tools.register(tags=["misc", "stock"])
def get_mock_stock_price(symbol: str) -> dict[str, Any]:
    """Get a mock stock price for a symbol."""
    return {
        "symbol": symbol.upper(),
        "price": round(random.uniform(10, 2000), 2),
        "change": round(random.uniform(-50, 50), 2),
        "currency": "USD",
        "timestamp": time.time()
    }

@tools.register(tags=["misc", "game"])
def roll_dice(sides: int = 6, count: int = 1) -> list[int]:
    """Roll a specified number of dice with a given number of sides."""
    return [random.randint(1, sides) for _ in range(count)]

@tools.register(tags=["misc", "game"])
def flip_coin() -> str:
    """Flip a coin and return 'Heads' or 'Tails'."""
    return random.choice(["Heads", "Tails"])

@tools.register(tags=["misc", "game"])
def generate_random_password_memorable(words_count: int = 4) -> str:
    """Generate a memorable password using random words."""
    words = ["apple", "banana", "cherry", "dragon", "eagle", "forest", "galaxy", "honey", "island", "jungle"]
    return "-".join(random.choices(words, k=words_count)) + str(random.randint(10, 99))

@tools.register(tags=["misc"])
def get_random_quote() -> str:
    """Get a random inspirational quote."""
    quotes = [
        "The only way to do great work is to love what you do. - Steve Jobs",
        "Innovation distinguishes between a leader and a follower. - Steve Jobs",
        "Your time is limited, so don't waste it living someone else's life. - Steve Jobs",
        "Stay hungry, stay foolish. - Steve Jobs",
        "The future belongs to those who believe in the beauty of their dreams. - Eleanor Roosevelt"
    ]
    return random.choice(quotes)

@tools.register(tags=["misc"])
def get_random_fact() -> str:
    """Get a random interesting fact."""
    facts = [
        "Honey never spoils.",
        "A day on Venus is longer than a year on Venus.",
        "Bananas are berries, but strawberries aren't.",
        "Octopuses have three hearts.",
        "A group of flamingos is called a 'flamboyance'."
    ]
    return random.choice(facts)

@tools.register(tags=["misc"])
def calculate_age(birth_date: str) -> int:
    """Calculate age based on birth date (YYYY-MM-DD)."""
    from datetime import datetime
    birth = datetime.strptime(birth_date, "%Y-%m-%d")
    today = datetime.now()
    return today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))

@tools.register(tags=["misc"])
def is_prime_advanced(n: int) -> bool:
    """Check if a number is prime using a more efficient algorithm."""
    if n <= 1: return False
    if n <= 3: return True
    if n % 2 == 0 or n % 3 == 0: return False
    i = 5
    while i * i <= n:
        if n % i == 0 or n % (i + 2) == 0: return False
        i += 6
    return True

@tools.register(tags=["misc"])
def get_fibonacci_sequence(n: int) -> list[int]:
    """Get the first n numbers of the Fibonacci sequence."""
    if n <= 0: return []
    if n == 1: return [0]
    seq = [0, 1]
    while len(seq) < n:
        seq.append(seq[-1] + seq[-2])
    return seq

@tools.register(tags=["misc"])
def solve_quadratic_equation(a: float, b: float, c: float) -> tuple[complex | None, complex | None]:
    """Solve a quadratic equation ax^2 + bx + c = 0."""
    import cmath
    d = (b**2) - (4*a*c)
    sol1 = (-b-cmath.sqrt(d))/(2*a)
    sol2 = (-b+cmath.sqrt(d))/(2*a)
    return sol1, sol2

@tools.register(tags=["misc"])
def get_lorem_ipsum(paragraphs: int = 1) -> str:
    """Generate mock Lorem Ipsum text."""
    text = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. " * 5
    return "\n\n".join([text] * paragraphs)

@tools.register(tags=["misc"])
def get_random_user_agent() -> str:
    """Get a random mock browser User-Agent string."""
    user_agents = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36",
        "Mozilla/5.0 (iPhone; CPU iPhone OS 17_1 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.1 Mobile/15E148 Safari/604.1"
    ]
    return random.choice(user_agents)

# --- Advanced Network & Security Tools ---

@tools.register(tags=["network", "security"])
def scan_ports(host: str, start_port: int, end_port: int) -> list[int]:
    """Scan a range of ports on a host and return a list of open ports."""
    open_ports = []
    for port in range(start_port, end_port + 1):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.settimeout(0.1)
            if s.connect_ex((host, port)) == 0:
                open_ports.append(port)
    return open_ports

@tools.register(tags=["network"])
def get_whois_info_mock(domain: str) -> dict[str, Any]:
    """Get mock WHOIS information for a domain."""
    return {
        "domain": domain,
        "registrar": "Mock Registrar LLC",
        "creation_date": "2010-01-01",
        "expiration_date": "2030-01-01",
        "name_servers": ["ns1.mock.com", "ns2.mock.com"]
    }

@tools.register(tags=["network"])
def get_public_ip() -> str:
    """Get the public IP address of the machine (requires internet)."""
    try:
        return requests.get('https://api.ipify.org', timeout=5).text
    except Exception:
        return "Unknown"

# --- Advanced Data & Math Tools ---

@tools.register(tags=["math"])
def calculate_standard_normal_cdf(x: float) -> float:
    """Calculate the Cumulative Distribution Function (CDF) of the standard normal distribution."""
    return (1.0 + math.erf(x / math.sqrt(2.0))) / 2.0

@tools.register(tags=["math"])
def calculate_entropy(probabilities: list[float]) -> float:
    """Calculate the Shannon entropy of a probability distribution."""
    return -sum(p * math.log2(p) for p in probabilities if p > 0)

@tools.register(tags=["data"])
def get_top_n_elements(elements: list[Any], n: int = 5) -> list[tuple[Any, int]]:
    """Get the top N most frequent elements in a list."""
    from collections import Counter
    return Counter(elements).most_common(n)

# --- Advanced Misc & Fun Tools ---

@tools.register(tags=["misc", "fun"])
def generate_ascii_art_text(text: str) -> str:
    """Generate simple ASCII art text (mock)."""
    return f"*** {text.upper()} ***"

@tools.register(tags=["misc", "fun"])
def get_random_emoji() -> str:
    """Get a random emoji."""
    emojis = ["😀", "🚀", "🔥", "🎉", "🤖", "🌈", "🍕", "🎸", "🌟", "🦄"]
    return random.choice(emojis)

@tools.register(tags=["misc", "fun"])
def get_random_color_name() -> str:
    """Get a random color name."""
    colors = ["Red", "Green", "Blue", "Yellow", "Magenta", "Cyan", "Orange", "Purple", "Pink", "Brown"]
    return random.choice(colors)

@tools.register(tags=["misc"])
def get_current_moon_phase_mock() -> str:
    """Get a mock current moon phase."""
    phases = ["New Moon", "Waxing Crescent", "First Quarter", "Waxing Gibbous", "Full Moon", "Waning Gibbous", "Last Quarter", "Waning Crescent"]
    return random.choice(phases)

@tools.register(tags=["misc"])
def get_random_zodiac_sign() -> str:
    """Get a random zodiac sign."""
    signs = ["Aries", "Taurus", "Gemini", "Cancer", "Leo", "Virgo", "Libra", "Scorpio", "Sagittarius", "Capricorn", "Aquarius", "Pisces"]
    return random.choice(signs)

@tools.register(tags=["misc"])
def get_random_planet() -> str:
    """Get a random planet in our solar system."""
    planets = ["Mercury", "Venus", "Earth", "Mars", "Jupiter", "Saturn", "Uranus", "Neptune"]
    return random.choice(planets)

@tools.register(tags=["misc"])
def get_random_chemical_element() -> dict[str, Any]:
    """Get a random chemical element."""
    elements = [
        {"name": "Hydrogen", "symbol": "H", "atomic_number": 1},
        {"name": "Helium", "symbol": "He", "atomic_number": 2},
        {"name": "Lithium", "symbol": "Li", "atomic_number": 3},
        {"name": "Carbon", "symbol": "C", "atomic_number": 6},
        {"name": "Oxygen", "symbol": "O", "atomic_number": 8},
        {"name": "Gold", "symbol": "Au", "atomic_number": 79}
    ]
    return random.choice(elements)

@tools.register(tags=["misc"])
def get_random_programming_language() -> str:
    """Get a random programming language."""
    languages = ["Python", "JavaScript", "Rust", "Go", "C++", "Java", "TypeScript", "Swift", "Kotlin", "Ruby"]
    return random.choice(languages)

@tools.register(tags=["misc"])
def get_random_http_status_code() -> int:
    """Get a random HTTP status code."""
    codes = [200, 201, 204, 400, 401, 403, 404, 500, 502, 503]
    return random.choice(codes)

@tools.register(tags=["misc"])
def get_random_uuid_v1() -> str:
    """Generate a random UUID (v1)."""
    return str(uuid.uuid1())

@tools.register(tags=["misc"])
def get_random_hex_string(length: int = 8) -> str:
    """Generate a random hexadecimal string."""
    return "".join(random.choices("0123456789abcdef", k=length))

@tools.register(tags=["misc"])
def get_random_base64_string(length: int = 12) -> str:
    """Generate a random Base64 string."""
    import secrets
    return secrets.token_urlsafe(length)

@tools.register(tags=["misc"])
def get_random_timestamp_recent() -> float:
    """Get a random timestamp from the last 24 hours."""
    return time.time() - random.uniform(0, 86400)

@tools.register(tags=["misc"])
def get_random_date_recent() -> str:
    """Get a random date string (YYYY-MM-DD) from the last year."""
    days_ago = random.randint(0, 365)
    d = datetime.date.today() - datetime.timedelta(days=days_ago)
    return d.isoformat()

@tools.register(tags=["misc"])
def get_random_time_recent() -> str:
    """Get a random time string (HH:MM:SS)."""
    return f"{random.randint(0, 23):02d}:{random.randint(0, 59):02d}:{random.randint(0, 59):02d}"

@tools.register(tags=["misc"])
def get_random_boolean_weighted(probability_true: float = 0.5) -> bool:
    """Generate a random boolean value with a specific probability of being True."""
    return random.random() < probability_true

@tools.register(tags=["misc"])
def get_random_float(min_val: float = 0.0, max_val: float = 1.0) -> float:
    """Generate a random float between min_val and max_val."""
    return random.uniform(min_val, max_val)

@tools.register(tags=["misc"])
def get_random_choice_weighted(choices: list[Any], weights: list[float]) -> Any:
    """Get a random choice from a list based on weights."""
    return random.choices(choices, weights=weights, k=1)[0]

@tools.register(tags=["misc"])
def get_random_sample(elements: list[Any], k: int) -> list[Any]:
    """Get a random sample of k elements from a list."""
    if k > len(elements): return elements
    return random.sample(elements, k)

@tools.register(tags=["misc"])
def get_random_gaussian(mu: float = 0.0, sigma: float = 1.0) -> float:
    """Generate a random number from a Gaussian distribution."""
    return random.gauss(mu, sigma)

@tools.register(tags=["misc"])
def get_random_exponential(lambd: float = 1.0) -> float:
    """Generate a random number from an exponential distribution."""
    return random.expovariate(lambd)

@tools.register(tags=["misc"])
def get_random_triangular(low: float, high: float, mode: float) -> float:
    """Generate a random number from a triangular distribution."""
    return random.triangular(low, high, mode)

@tools.register(tags=["misc"])
def get_random_beta(alpha: float, beta: float) -> float:
    """Generate a random number from a beta distribution."""
    return random.betavariate(alpha, beta)

@tools.register(tags=["misc"])
def get_random_gamma(alpha: float, beta: float) -> float:
    """Generate a random number from a gamma distribution."""
    return random.gammavariate(alpha, beta)

@tools.register(tags=["misc"])
def get_random_lognormal(mu: float, sigma: float) -> float:
    """Generate a random number from a log-normal distribution."""
    return random.lognormvariate(mu, sigma)

@tools.register(tags=["misc"])
def get_random_vonmises(mu: float, kappa: float) -> float:
    """Generate a random number from a Von Mises distribution."""
    return random.vonmisesvariate(mu, kappa)

@tools.register(tags=["misc"])
def get_random_pareto(alpha: float) -> float:
    """Generate a random number from a Pareto distribution."""
    return random.paretovariate(alpha)

@tools.register(tags=["misc"])
def get_random_weibull(alpha: float, beta: float) -> float:
    """Generate a random number from a Weibull distribution."""
    return random.weibullvariate(alpha, beta)

# --- Final Batch of Advanced Tools ---

@tools.register(tags=["data", "json"])
def validate_json_schema(data: dict[str, Any], schema: dict[str, Any]) -> bool:
    """Validate a dictionary against a JSON schema (requires jsonschema)."""
    try:
        import jsonschema
        jsonschema.validate(instance=data, schema=schema)
        return True
    except Exception:
        return False

@tools.register(tags=["web", "api"])
def get_github_repo_info_mock(repo_name: str) -> dict[str, Any]:
    """Get mock information about a GitHub repository."""
    return {
        "full_name": repo_name,
        "stars": random.randint(100, 100000),
        "forks": random.randint(10, 10000),
        "open_issues": random.randint(0, 500),
        "language": "Python",
        "license": "MIT"
    }

@tools.register(tags=["web", "api"])
def get_mock_crypto_price(coin: str) -> dict[str, Any]:
    """Get a mock cryptocurrency price."""
    return {
        "coin": coin.upper(),
        "price_usd": round(random.uniform(0.1, 60000), 2),
        "change_24h": round(random.uniform(-10, 10), 2),
        "market_cap": random.randint(1000000, 1000000000000)
    }

@tools.register(tags=["system"])
def get_system_uptime_mock() -> str:
    """Get a mock system uptime string."""
    days = random.randint(0, 30)
    hours = random.randint(0, 23)
    minutes = random.randint(0, 59)
    return f"{days} days, {hours} hours, {minutes} minutes"

@tools.register(tags=["text"])
def generate_slug(text: str) -> str:
    """Generate a URL-friendly slug from a string."""
    text = text.lower()
    text = re.sub(r'[^\w\s-]', '', text)
    return re.sub(r'[-\s]+', '-', text).strip('-')

@tools.register(tags=["math"])
def calculate_haversine_distance(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    """Calculate the great-circle distance between two points on Earth in km."""
    R = 6371.0  # Earth radius in km
    dlat = math.radians(lat2 - lat1)
    dlon = math.radians(lon2 - lon1)
    a = math.sin(dlat / 2)**2 + math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.sin(dlon / 2)**2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    return R * c

@tools.register(tags=["misc"])
def get_random_hex_color_vibrant() -> str:
    """Generate a random vibrant hex color."""
    h = random.random()
    s = 0.8
    v = 0.9
    import colorsys
    r, g, b = colorsys.hsv_to_rgb(h, s, v)
    return f"#{int(r*255):02x}{int(g*255):02x}{int(b*255):02x}"

@tools.register(tags=["data"])
def chunk_list(elements: list[Any], chunk_size: int) -> list[list[Any]]:
    """Split a list into chunks of a specified size."""
    return [elements[i:i + chunk_size] for i in range(0, len(elements), chunk_size)]

@tools.register(tags=["text"])
def wrap_text(text: str, width: int = 80) -> str:
    """Wrap text to a specified width."""
    import textwrap
    return textwrap.fill(text, width=width)

@tools.register(tags=["system"])
def get_python_path() -> list[str]:
    """Get the current Python system path."""
    return sys.path

@tools.register(tags=["misc"])
def get_random_mac_address_multicast() -> str:
    """Generate a random multicast MAC address."""
    first_byte = random.randint(0, 255) | 0x01
    return f"{first_byte:02x}:" + ":".join(f"{random.randint(0, 255):02x}" for _ in range(5))

@tools.register(tags=["math"])
def calculate_log10(x: float) -> float:
    """Calculate the base-10 logarithm of x."""
    return math.log10(x)

@tools.register(tags=["math"])
def calculate_exp(x: float) -> float:
    """Calculate the exponential of x (e^x)."""
    return math.exp(x)

@tools.register(tags=["text"])
def is_upper(text: str) -> bool:
    """Check if a string is all uppercase."""
    return text.isupper()

@tools.register(tags=["text"])
def is_lower(text: str) -> bool:
    """Check if a string is all lowercase."""
    return text.islower()

@tools.register(tags=["data"])
def get_dict_keys(d: dict[str, Any]) -> list[str]:
    """Get all keys of a dictionary."""
    return list(d.keys())

@tools.register(tags=["data"])
def get_dict_values(d: dict[str, Any]) -> list[Any]:
    """Get all values of a dictionary."""
    return list(d.values())

@tools.register(tags=["misc"])
def get_random_int_range(start: int, stop: int, step: int = 1) -> int:
    """Generate a random integer from range(start, stop, step)."""
    return random.randrange(start, stop, step)

@tools.register(tags=["misc"])
def get_random_bits(k: int) -> int:
    """Generate an integer with k random bits."""
    return random.getrandbits(k)

@tools.register(tags=["system"])
def get_current_process_name() -> str:
    """Get the name of the current process."""
    import multiprocessing
    return multiprocessing.current_process().name

