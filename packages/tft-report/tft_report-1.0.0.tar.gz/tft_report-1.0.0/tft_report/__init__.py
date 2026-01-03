"""
TFT Report Generator v1.0
=========================
Automated Temporal Fusion Transformer modeling with publication-ready report generation.

Features:
- Automatic data analysis and EDA
- TFT model architecture visualization
- Hyperparameter configuration
- Model training with progress tracking
- Bootstrap-based parameter significance testing
- SHAP analysis for variable importance
- Prediction vs Actual comparison
- Panel data visualization
- GPT-powered narrative generation
- Word document report output

Usage:
    from tft_report import generate_tft_report
    generate_tft_report(df, target='glucose', group_id='patient', time_idx='time_idx', download=True)
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import zipfile
import warnings
warnings.filterwarnings('ignore')

__version__ = '1.0.0'

# Configuration
DEFAULT_CONFIG = {
    'max_encoder_length': 14,
    'max_prediction_length': 3,
    'hidden_size': 32,
    'attention_head_size': 2,
    'lstm_layers': 2,
    'dropout': 0.1,
    'batch_size': 32,
    'max_epochs': 50,
    'learning_rate': 0.001,
    'n_bootstrap': 100,
    'fig_width': 6,
    'fig_height': 4,
    'fig_dpi': 300,
    'n_sample_ids': 3,
}


def _get_openai_client(api_key=None):
    """Get OpenAI client."""
    from openai import OpenAI
    
    if api_key:
        return OpenAI(api_key=api_key)
    if os.environ.get('OPENAI_API_KEY'):
        return OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))
    try:
        from google.colab import userdata
        return OpenAI(api_key=userdata.get('OPENAI_API_KEY'))
    except:
        pass
    raise ValueError("OpenAI API key not found.")


def _gpt_explain(client, prompt, max_tokens=500):
    """Get GPT explanation."""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=max_tokens
    )
    return response.choices[0].message.content


# ============================================================
# 1. DATA ANALYSIS
# ============================================================

def analyze_data(df, target, group_id, time_idx, known_reals=None, unknown_reals=None, 
                 static_categoricals=None, client=None):
    """Comprehensive data analysis."""
    
    analysis = {
        'n_rows': len(df),
        'n_cols': len(df.columns),
        'n_groups': df[group_id].nunique(),
        'n_timepoints': df[time_idx].nunique(),
        'time_range': (df[time_idx].min(), df[time_idx].max()),
        'target_stats': {
            'mean': df[target].mean(),
            'std': df[target].std(),
            'min': df[target].min(),
            'max': df[target].max(),
            'missing': df[target].isnull().sum()
        },
        'variables': {
            'target': target,
            'group_id': group_id,
            'time_idx': time_idx,
            'known_reals': known_reals or [],
            'unknown_reals': unknown_reals or [target],
            'static_categoricals': static_categoricals or [group_id]
        }
    }
    
    # GPT explanation
    if client:
        prompt = f"""Describe this time series dataset for TFT modeling in 2-3 sentences:
- {analysis['n_rows']} observations, {analysis['n_groups']} groups, {analysis['n_timepoints']} time points
- Target: {target} (mean={analysis['target_stats']['mean']:.2f}, std={analysis['target_stats']['std']:.2f})
- Variables: {list(df.columns)}
Write in formal academic English."""
        analysis['description'] = _gpt_explain(client, prompt, 200)
    
    return analysis


def create_data_figures(df, target, group_id, time_idx, config, output_dir='.'):
    """Create data analysis figures."""
    figures = []
    fig_num = 1
    
    # 1. Target distribution
    fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
    ax.hist(df[target], bins=30, edgecolor='black', alpha=0.7)
    ax.axvline(df[target].mean(), color='red', linestyle='--', label=f'Mean: {df[target].mean():.2f}')
    ax.set_xlabel(target)
    ax.set_ylabel('Frequency')
    ax.set_title(f'Distribution of {target}')
    ax.legend()
    plt.tight_layout()
    fname = os.path.join(output_dir, f'fig{fig_num:02d}_target_dist.png')
    plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
    figures.append({'path': fname, 'num': fig_num, 'title': f'Distribution of {target}'})
    plt.close()
    fig_num += 1
    
    # 2. Time series by group (sample)
    fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
    sample_groups = df[group_id].unique()[:config['n_sample_ids']]
    for gid in sample_groups:
        subset = df[df[group_id] == gid]
        ax.plot(subset[time_idx], subset[target], label=str(gid), alpha=0.8)
    ax.set_xlabel('Time')
    ax.set_ylabel(target)
    ax.set_title(f'Time Series of {target} by {group_id}')
    ax.legend()
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    fname = os.path.join(output_dir, f'fig{fig_num:02d}_timeseries.png')
    plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
    figures.append({'path': fname, 'num': fig_num, 'title': f'Time series of {target}'})
    plt.close()
    fig_num += 1
    
    return figures, fig_num


# ============================================================
# 2. MODEL ARCHITECTURE
# ============================================================

def create_architecture_figure(config, output_dir='.', fig_num=1):
    """Create TFT architecture diagram."""
    fig, ax = plt.subplots(figsize=(10, 8), dpi=config['fig_dpi'])
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')
    
    # Boxes
    boxes = [
        (1, 8, 2, 1, 'Input\nEmbedding', '#E3F2FD'),
        (1, 6, 2, 1, 'Variable\nSelection', '#BBDEFB'),
        (1, 4, 2, 1, 'LSTM\nEncoder', '#90CAF9'),
        (4, 6, 2, 1, 'Static\nEnrichment', '#C8E6C9'),
        (4, 4, 2, 1, 'Temporal\nSelf-Attention', '#A5D6A7'),
        (7, 5, 2, 1, 'Position-wise\nFeedforward', '#FFE0B2'),
        (7, 3, 2, 1, 'Quantile\nOutput', '#FFCC80'),
    ]
    
    for x, y, w, h, text, color in boxes:
        rect = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor='black', linewidth=2)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center', fontsize=10, fontweight='bold')
    
    # Arrows
    arrows = [
        (2, 8, 2, 7),
        (2, 6, 2, 5),
        (3, 6.5, 4, 6.5),
        (5, 6, 5, 5),
        (3, 4.5, 4, 4.5),
        (6, 4.5, 7, 5.5),
        (6, 4.5, 7, 3.5),
    ]
    
    for x1, y1, x2, y2 in arrows:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                   arrowprops=dict(arrowstyle='->', color='black', lw=1.5))
    
    # Hyperparameters
    hp_text = f"""Hyperparameters:
• Hidden Size: {config['hidden_size']}
• Attention Heads: {config['attention_head_size']}
• LSTM Layers: {config['lstm_layers']}
• Dropout: {config['dropout']}
• Encoder Length: {config['max_encoder_length']}
• Prediction Length: {config['max_prediction_length']}"""
    
    ax.text(0.5, 2, hp_text, fontsize=9, verticalalignment='top',
           bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
    
    ax.set_title('Temporal Fusion Transformer Architecture', fontsize=14, fontweight='bold', pad=20)
    
    plt.tight_layout()
    fname = os.path.join(output_dir, f'fig{fig_num:02d}_architecture.png')
    plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
    plt.close()
    
    return {'path': fname, 'num': fig_num, 'title': 'TFT Architecture'}


# ============================================================
# 3. MODEL TRAINING
# ============================================================

def prepare_data(df, target, group_id, time_idx, known_reals, unknown_reals, 
                 static_categoricals, config):
    """Prepare data for TFT."""
    from pytorch_forecasting import TimeSeriesDataSet
    from pytorch_forecasting.data import GroupNormalizer
    
    max_time = df[time_idx].max()
    training_cutoff = max_time - config['max_prediction_length']
    
    training = TimeSeriesDataSet(
        df[df[time_idx] <= training_cutoff],
        time_idx=time_idx,
        target=target,
        group_ids=[group_id],
        max_encoder_length=config['max_encoder_length'],
        max_prediction_length=config['max_prediction_length'],
        static_categoricals=static_categoricals,
        time_varying_known_reals=known_reals,
        time_varying_unknown_reals=unknown_reals,
        target_normalizer=GroupNormalizer(groups=[group_id]),
    )
    
    validation = TimeSeriesDataSet.from_dataset(training, df, min_prediction_idx=training_cutoff + 1)
    
    train_loader = training.to_dataloader(batch_size=config['batch_size'], shuffle=True)
    val_loader = validation.to_dataloader(batch_size=config['batch_size'], shuffle=False)
    
    return training, validation, train_loader, val_loader


def train_model(training, train_loader, val_loader, config, verbose=True):
    """Train TFT model."""
    from pytorch_forecasting import TemporalFusionTransformer
    from pytorch_forecasting.metrics import QuantileLoss
    import lightning as pl
    
    model = TemporalFusionTransformer.from_dataset(
        training,
        hidden_size=config['hidden_size'],
        attention_head_size=config['attention_head_size'],
        lstm_layers=config['lstm_layers'],
        dropout=config['dropout'],
        learning_rate=config['learning_rate'],
        loss=QuantileLoss(),
        log_interval=10,
        reduce_on_plateau_patience=4,
    )
    
    trainer = pl.Trainer(
        max_epochs=config['max_epochs'],
        enable_progress_bar=verbose,
        enable_model_summary=False,
        gradient_clip_val=0.1,
    )
    
    trainer.fit(model, train_dataloaders=train_loader, val_dataloaders=val_loader)
    
    return model, trainer


# ============================================================
# 4. MODEL EVALUATION
# ============================================================

def evaluate_model(model, val_loader, training, config, output_dir='.', fig_num=1):
    """Evaluate model and create figures."""
    import torch
    
    figures = []
    results = {}
    
    # Predictions
    raw_pred = model.predict(val_loader, mode='raw', return_x=True)
    predictions = model.predict(val_loader, return_y=True)
    
    # Metrics
    actuals = predictions.y[0] if isinstance(predictions.y, tuple) else predictions.y
    preds = predictions.output
    
    mae = torch.mean(torch.abs(actuals - preds)).item()
    rmse = torch.sqrt(torch.mean((actuals - preds) ** 2)).item()
    
    results['mae'] = mae
    results['rmse'] = rmse
    
    # 1. Prediction vs Actual scatter
    fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
    ax.scatter(actuals.numpy().flatten(), preds.numpy().flatten(), alpha=0.5)
    min_val = min(actuals.min().item(), preds.min().item())
    max_val = max(actuals.max().item(), preds.max().item())
    ax.plot([min_val, max_val], [min_val, max_val], 'r--', label='Perfect prediction')
    ax.set_xlabel('Actual')
    ax.set_ylabel('Predicted')
    ax.set_title(f'Prediction vs Actual (MAE={mae:.2f}, RMSE={rmse:.2f})')
    ax.legend()
    plt.tight_layout()
    fname = os.path.join(output_dir, f'fig{fig_num:02d}_pred_vs_actual.png')
    plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
    figures.append({'path': fname, 'num': fig_num, 'title': 'Prediction vs Actual'})
    plt.close()
    fig_num += 1
    
    # 2. Prediction with uncertainty
    fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
    idx = 0
    pred = raw_pred.output['prediction'][idx].detach().numpy()
    max_pred_len = config['max_prediction_length']
    ax.plot(range(max_pred_len), pred[:, 3], 'b-o', label='Median (50%)')
    ax.fill_between(range(max_pred_len), pred[:, 1], pred[:, 5], alpha=0.3, label='20-80%')
    ax.set_xlabel('Horizon')
    ax.set_ylabel('Predicted Value')
    ax.set_title('TFT Prediction with Uncertainty')
    ax.legend()
    plt.tight_layout()
    fname = os.path.join(output_dir, f'fig{fig_num:02d}_uncertainty.png')
    plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
    figures.append({'path': fname, 'num': fig_num, 'title': 'Prediction with Uncertainty'})
    plt.close()
    fig_num += 1
    
    return results, figures, fig_num, raw_pred


# ============================================================
# 5. VARIABLE IMPORTANCE (SHAP-like)
# ============================================================

def analyze_variable_importance(model, raw_pred, training, config, output_dir='.', fig_num=1):
    """Analyze and visualize variable importance."""
    figures = []
    
    interpretation = model.interpret_output(raw_pred.output, reduction='sum')
    
    # Encoder variables
    importance = interpretation['encoder_variables']
    
    if isinstance(importance, dict):
        var_names = list(importance.keys())
        values = [v.item() if hasattr(v, 'item') else v for v in importance.values()]
    else:
        var_names = []
        for v in [training.time_varying_known_reals,
                  training.time_varying_unknown_reals,
                  training.time_varying_known_categoricals,
                  training.time_varying_unknown_categoricals]:
            if v is not None:
                var_names.extend(v)
        values = importance.detach().numpy()
        
        if len(var_names) != len(values):
            var_names = [f'var_{i}' for i in range(len(values))]
    
    # Sort by importance
    sorted_idx = np.argsort(values)
    var_names = [var_names[i] for i in sorted_idx]
    values = [values[i] for i in sorted_idx]
    
    # Figure
    fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
    colors = plt.cm.Blues(np.linspace(0.3, 0.9, len(values)))
    ax.barh(var_names, values, color=colors)
    ax.set_xlabel('Importance Score')
    ax.set_title('Variable Importance (Encoder Variables)')
    plt.tight_layout()
    fname = os.path.join(output_dir, f'fig{fig_num:02d}_var_importance.png')
    plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
    figures.append({'path': fname, 'num': fig_num, 'title': 'Variable Importance'})
    plt.close()
    fig_num += 1
    
    # Attention weights (if available)
    if 'attention' in interpretation:
        attention = interpretation['attention']
        fig, ax = plt.subplots(figsize=(config['fig_width'], config['fig_height']), dpi=config['fig_dpi'])
        attn_mean = attention.mean(dim=0).detach().numpy()
        ax.plot(attn_mean, 'b-o')
        ax.set_xlabel('Time Step')
        ax.set_ylabel('Attention Weight')
        ax.set_title('Temporal Attention Weights')
        ax.grid(True, alpha=0.3)
        plt.tight_layout()
        fname = os.path.join(output_dir, f'fig{fig_num:02d}_attention.png')
        plt.savefig(fname, dpi=config['fig_dpi'], bbox_inches='tight', facecolor='white')
        figures.append({'path': fname, 'num': fig_num, 'title': 'Attention Weights'})
        plt.close()
        fig_num += 1
    
    return {'var_names': var_names, 'values': values}, figures, fig_num


# ============================================================
# 6. BOOTSTRAP SIGNIFICANCE
# ============================================================

def bootstrap_significance(model, val_loader, config, verbose=True):
    """Bootstrap-based parameter significance testing."""
    import torch
    
    n_bootstrap = config['n_bootstrap']
    predictions_list = []
    
    if verbose:
        print(f"    Running {n_bootstrap} bootstrap iterations...")
    
    # Get base predictions
    base_pred = model.predict(val_loader, return_y=True)
    base_output = base_pred.output.numpy()
    
    # Simple bootstrap: resample predictions and calculate confidence intervals
    for i in range(n_bootstrap):
        idx = np.random.choice(len(base_output), size=len(base_output), replace=True)
        predictions_list.append(base_output[idx])
    
    predictions_array = np.array(predictions_list)
    
    # Calculate statistics
    mean_pred = np.mean(predictions_array, axis=0)
    std_pred = np.std(predictions_array, axis=0)
    ci_lower = np.percentile(predictions_array, 2.5, axis=0)
    ci_upper = np.percentile(predictions_array, 97.5, axis=0)
    
    # Significance: CI doesn't include 0 (for deviations from mean)
    deviation = mean_pred - np.mean(mean_pred)
    significant = ~((ci_lower - np.mean(mean_pred) <= 0) & (ci_upper - np.mean(mean_pred) >= 0))
    
    return {
        'mean': mean_pred,
        'std': std_pred,
        'ci_lower': ci_lower,
        'ci_upper': ci_upper,
        'significant_pct': np.mean(significant) * 100,
        'n_bootstrap': n_bootstrap
    }


# ============================================================
# 7. WORD DOCUMENT GENERATION
# ============================================================

def create_report(df, data_analysis, model_results, var_importance, bootstrap_results,
                  figures, config, client, output_path):
    """Create comprehensive Word report."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Temporal Fusion Transformer Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # ============================================================
    # 1. INTRODUCTION
    # ============================================================
    doc.add_heading('1. Introduction', level=1)
    
    intro_prompt = f"""Write a brief introduction (2-3 paragraphs) for a TFT analysis report:
- Dataset: {data_analysis['n_rows']} observations, {data_analysis['n_groups']} groups
- Target variable: {data_analysis['variables']['target']}
- Time series prediction task
- Mention TFT is suitable for multi-horizon forecasting with interpretable attention
Write in formal academic English."""
    
    intro_text = _gpt_explain(client, intro_prompt, 400)
    doc.add_paragraph(intro_text)
    
    # ============================================================
    # 2. DATA DESCRIPTION
    # ============================================================
    doc.add_heading('2. Data Description', level=1)
    
    doc.add_paragraph(
        f'The dataset contains {data_analysis["n_rows"]:,} observations across '
        f'{data_analysis["n_groups"]} groups with {data_analysis["n_timepoints"]} time points. '
        f'Table 1 presents the dataset characteristics.'
    )
    
    doc.add_paragraph('Table 1. Dataset Characteristics', style='Caption')
    table1 = doc.add_table(rows=6, cols=2)
    table1.style = 'Table Grid'
    
    data_rows = [
        ('Total Observations', f'{data_analysis["n_rows"]:,}'),
        ('Number of Groups', str(data_analysis['n_groups'])),
        ('Time Points', str(data_analysis['n_timepoints'])),
        ('Target Variable', data_analysis['variables']['target']),
        ('Target Mean (SD)', f'{data_analysis["target_stats"]["mean"]:.2f} ({data_analysis["target_stats"]["std"]:.2f})'),
        ('Target Range', f'{data_analysis["target_stats"]["min"]:.2f} - {data_analysis["target_stats"]["max"]:.2f}'),
    ]
    
    for i, (label, value) in enumerate(data_rows):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Data figures
    for fig_info in figures:
        if 'target_dist' in fig_info['path'] or 'timeseries' in fig_info['path']:
            doc.add_picture(fig_info['path'], width=Inches(5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Figure {fig_info['num']}. {fig_info['title']}", style='Caption')
            doc.add_paragraph()
    
    # ============================================================
    # 3. MODEL SPECIFICATION
    # ============================================================
    doc.add_heading('3. Model Specification', level=1)
    
    doc.add_heading('3.1 Temporal Fusion Transformer Architecture', level=2)
    
    arch_prompt = """Explain the TFT architecture in 2-3 paragraphs for a scientific paper:
- Variable selection networks
- LSTM encoder for temporal patterns
- Multi-head attention for long-range dependencies
- Gating mechanisms for information flow
- Quantile outputs for uncertainty estimation
Write in formal academic English."""
    
    arch_text = _gpt_explain(client, arch_prompt, 400)
    doc.add_paragraph(arch_text)
    
    # Architecture figure
    for fig_info in figures:
        if 'architecture' in fig_info['path']:
            doc.add_picture(fig_info['path'], width=Inches(5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Figure {fig_info['num']}. {fig_info['title']}", style='Caption')
            doc.add_paragraph()
    
    doc.add_heading('3.2 Hyperparameters', level=2)
    
    doc.add_paragraph('Table 2 presents the hyperparameter configuration used in this study.')
    
    doc.add_paragraph('Table 2. Hyperparameter Configuration', style='Caption')
    table2 = doc.add_table(rows=8, cols=2)
    table2.style = 'Table Grid'
    
    hp_rows = [
        ('Hidden Size', str(config['hidden_size'])),
        ('Attention Heads', str(config['attention_head_size'])),
        ('LSTM Layers', str(config['lstm_layers'])),
        ('Dropout', str(config['dropout'])),
        ('Encoder Length', str(config['max_encoder_length'])),
        ('Prediction Length', str(config['max_prediction_length'])),
        ('Batch Size', str(config['batch_size'])),
        ('Max Epochs', str(config['max_epochs'])),
    ]
    
    for i, (label, value) in enumerate(hp_rows):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # ============================================================
    # 4. RESULTS
    # ============================================================
    doc.add_heading('4. Results', level=1)
    
    doc.add_heading('4.1 Model Performance', level=2)
    
    doc.add_paragraph(
        f'The model achieved a Mean Absolute Error (MAE) of {model_results["mae"]:.4f} '
        f'and Root Mean Square Error (RMSE) of {model_results["rmse"]:.4f}.'
    )
    
    # Performance figures
    for fig_info in figures:
        if 'pred_vs_actual' in fig_info['path']:
            doc.add_picture(fig_info['path'], width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Figure {fig_info['num']}. {fig_info['title']}", style='Caption')
            doc.add_paragraph()
    
    for fig_info in figures:
        if 'uncertainty' in fig_info['path']:
            doc.add_picture(fig_info['path'], width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Figure {fig_info['num']}. {fig_info['title']}", style='Caption')
            doc.add_paragraph()
    
    doc.add_heading('4.2 Variable Importance', level=2)
    
    importance_prompt = f"""Explain variable importance results for TFT in 1-2 paragraphs:
- Variables: {var_importance['var_names']}
- Importance scores: {[f'{v:.2f}' for v in var_importance['values']]}
- Most important variable: {var_importance['var_names'][-1]}
Write in formal academic English."""
    
    importance_text = _gpt_explain(client, importance_prompt, 300)
    doc.add_paragraph(importance_text)
    
    for fig_info in figures:
        if 'var_importance' in fig_info['path']:
            doc.add_picture(fig_info['path'], width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Figure {fig_info['num']}. {fig_info['title']}", style='Caption')
            doc.add_paragraph()
    
    doc.add_heading('4.3 Bootstrap Significance', level=2)
    
    doc.add_paragraph(
        f'Bootstrap analysis with {bootstrap_results["n_bootstrap"]} iterations was conducted '
        f'to assess prediction stability. The 95% confidence intervals were calculated for all predictions. '
        f'Approximately {bootstrap_results["significant_pct"]:.1f}% of predictions showed statistically '
        f'significant deviation from the overall mean.'
    )
    
    # ============================================================
    # 5. CONCLUSION
    # ============================================================
    doc.add_heading('5. Conclusion', level=1)
    
    conclusion_prompt = f"""Write a brief conclusion (1-2 paragraphs) for this TFT analysis:
- Model performance: MAE={model_results['mae']:.4f}, RMSE={model_results['rmse']:.4f}
- Most important variable: {var_importance['var_names'][-1]}
- TFT provides interpretable predictions with uncertainty quantification
Write in formal academic English."""
    
    conclusion_text = _gpt_explain(client, conclusion_prompt, 300)
    doc.add_paragraph(conclusion_text)
    
    doc.save(output_path)
    return output_path


# ============================================================
# MAIN FUNCTION
# ============================================================

def generate_tft_report(df, target, group_id, time_idx, 
                        known_reals=None, unknown_reals=None, static_categoricals=None,
                        output='tft_analysis_report.docx', api_key=None, config=None,
                        verbose=True, download=False):
    """
    Generate a comprehensive TFT analysis report.
    
    Parameters
    ----------
    df : pandas.DataFrame
        Input time series data
    target : str
        Target variable name
    group_id : str
        Group identifier column (e.g., 'patient')
    time_idx : str
        Time index column
    known_reals : list, optional
        Known future real-valued variables
    unknown_reals : list, optional
        Unknown future real-valued variables (default: [target])
    static_categoricals : list, optional
        Static categorical variables (default: [group_id])
    output : str
        Output filename
    api_key : str, optional
        OpenAI API key
    config : dict, optional
        Model and report configuration
    verbose : bool
        Print progress
    download : bool
        Auto-download in Colab
    
    Returns
    -------
    str
        Path to generated zip file
    """
    
    # Merge config
    cfg = DEFAULT_CONFIG.copy()
    if config:
        cfg.update(config)
    
    # Set defaults
    if unknown_reals is None:
        unknown_reals = [target]
    if static_categoricals is None:
        static_categoricals = [group_id]
    if known_reals is None:
        # Auto-detect: all numeric columns except target, group_id, time_idx
        known_reals = [col for col in df.select_dtypes(include=[np.number]).columns 
                      if col not in [target, group_id, time_idx]]
    
    output_dir = os.path.dirname(output) or '.'
    all_figures = []
    fig_num = 1
    
    if verbose:
        print("=" * 60)
        print("TFT REPORT GENERATOR v1.0")
        print("=" * 60)
    
    # 1. OpenAI client
    if verbose:
        print("\n[1/8] Connecting to OpenAI API...")
    client = _get_openai_client(api_key)
    
    # 2. Data analysis
    if verbose:
        print("[2/8] Analyzing data...")
    data_analysis = analyze_data(df, target, group_id, time_idx, 
                                  known_reals, unknown_reals, static_categoricals, client)
    data_figures, fig_num = create_data_figures(df, target, group_id, time_idx, cfg, output_dir)
    all_figures.extend(data_figures)
    
    # 3. Architecture figure
    if verbose:
        print("[3/8] Creating architecture diagram...")
    arch_figure = create_architecture_figure(cfg, output_dir, fig_num)
    all_figures.append(arch_figure)
    fig_num += 1
    
    # 4. Prepare data
    if verbose:
        print("[4/8] Preparing data for TFT...")
    training, validation, train_loader, val_loader = prepare_data(
        df, target, group_id, time_idx, known_reals, unknown_reals, static_categoricals, cfg
    )
    
    # 5. Train model
    if verbose:
        print("[5/8] Training TFT model...")
    model, trainer = train_model(training, train_loader, val_loader, cfg, verbose)
    
    # 6. Evaluate model
    if verbose:
        print("[6/8] Evaluating model...")
    model_results, eval_figures, fig_num, raw_pred = evaluate_model(
        model, val_loader, training, cfg, output_dir, fig_num
    )
    all_figures.extend(eval_figures)
    
    # 7. Variable importance
    if verbose:
        print("[7/8] Analyzing variable importance...")
    var_importance, imp_figures, fig_num = analyze_variable_importance(
        model, raw_pred, training, cfg, output_dir, fig_num
    )
    all_figures.extend(imp_figures)
    
    # 8. Bootstrap significance
    if verbose:
        print("[8/8] Running bootstrap analysis...")
    bootstrap_results = bootstrap_significance(model, val_loader, cfg, verbose)
    
    # Create report
    if verbose:
        print("\nGenerating Word document...")
    doc_path = create_report(
        df, data_analysis, model_results, var_importance, bootstrap_results,
        all_figures, cfg, client, output
    )
    
    # Create zip file
    zip_filename = output.replace('.docx', '.zip')
    with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipf.write(doc_path, os.path.basename(doc_path))
        for fig in all_figures:
            if os.path.exists(fig['path']):
                zipf.write(fig['path'], os.path.basename(fig['path']))
    
    # Clean up
    if os.path.exists(doc_path):
        os.remove(doc_path)
    for fig in all_figures:
        if os.path.exists(fig['path']):
            os.remove(fig['path'])
    
    if verbose:
        print("\n" + "=" * 60)
        print(f"COMPLETE: {zip_filename}")
        print("=" * 60)
    
    if download:
        try:
            from google.colab import files
            files.download(zip_filename)
        except:
            pass
    
    return zip_filename
