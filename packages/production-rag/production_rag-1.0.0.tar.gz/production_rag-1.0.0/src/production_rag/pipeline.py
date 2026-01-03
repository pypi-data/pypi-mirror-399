"""
RAG Pipeline - Main orchestration for RAG operations.
"""

import time
import logging
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, field
from pathlib import Path

from .types import Document, Chunk, Query, Response, RetrievalResult, DocumentType
from .vectorstore import VectorStore, VectorStoreConfig
from .retriever import Retriever, RetrieverConfig
from .chunker import Chunker
from .embedder import Embedder, EmbeddingConfig
from .reranker import Reranker
from .generator import Generator, GeneratorConfig
from .cache import RAGCache, CacheConfig

logger = logging.getLogger(__name__)


@dataclass
class PipelineConfig:
    """Configuration for RAG Pipeline."""
    
    # Chunking
    chunk_size: int = 512
    chunk_overlap: int = 50
    chunking_strategy: str = "recursive"
    
    # Embedding
    embedding_model: str = "all-MiniLM-L6-v2"
    embedding_dim: int = 384
    
    # Retrieval
    top_k: int = 5
    similarity_threshold: float = 0.7
    use_reranker: bool = True
    reranker_model: str = "cross-encoder/ms-marco-MiniLM-L-6-v2"
    
    # Generation
    llm_model: str = "gpt-3.5-turbo"
    max_tokens: int = 1024
    temperature: float = 0.7
    
    # Caching
    enable_cache: bool = True
    cache_ttl: int = 3600
    
    # Performance
    batch_size: int = 32
    max_concurrent: int = 4
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "embedding_model": self.embedding_model,
            "top_k": self.top_k,
            "use_reranker": self.use_reranker,
            "llm_model": self.llm_model,
            "enable_cache": self.enable_cache,
        }


class RAGPipeline:
    """
    Production-ready RAG Pipeline.
    
    Orchestrates the complete RAG workflow:
    1. Document ingestion and chunking
    2. Embedding generation
    3. Vector storage and indexing
    4. Query processing and retrieval
    5. Context reranking
    6. Response generation
    
    Examples:
        Basic usage:
        >>> pipeline = RAGPipeline()
        >>> pipeline.add_documents(["doc.pdf", "doc.txt"])
        >>> response = pipeline.query("What is the main topic?")
        >>> print(response.answer)
        
        With configuration:
        >>> config = PipelineConfig(chunk_size=1000, top_k=10)
        >>> pipeline = RAGPipeline(config=config)
        
        With custom components:
        >>> pipeline = RAGPipeline()
        >>> pipeline.set_generator(CustomGenerator())
    """
    
    def __init__(
        self,
        config: Optional[PipelineConfig] = None,
        api_key: Optional[str] = None,
    ):
        """
        Initialize RAG Pipeline.
        
        Args:
            config: Pipeline configuration
            api_key: API key for LLM/embedding services
        """
        self.config = config or PipelineConfig()
        self.api_key = api_key
        
        # Initialize components
        self._init_components()
        
        # Statistics
        self._query_count = 0
        self._total_retrieval_time = 0.0
        self._total_generation_time = 0.0
        
        logger.info(f"RAGPipeline initialized with config: {self.config.to_dict()}")
    
    def _init_components(self) -> None:
        """Initialize pipeline components."""
        # Chunker
        self.chunker = Chunker(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            strategy=self.config.chunking_strategy,
        )
        
        # Embedder
        self.embedder = Embedder(
            model_name=self.config.embedding_model,
        )
        
        # Vector Store
        self.vector_store = VectorStore(
            embedding_dim=self.config.embedding_dim,
        )
        
        # Retriever
        self.retriever = Retriever(
            vector_store=self.vector_store,
            embedder=self.embedder,
            config=RetrieverConfig(
                top_k=self.config.top_k,
                similarity_threshold=self.config.similarity_threshold,
            ),
        )
        
        # Reranker
        if self.config.use_reranker:
            self.reranker = Reranker(
                model_name=self.config.reranker_model,
            )
        else:
            self.reranker = None
        
        # Generator
        self.generator = Generator(
            config=GeneratorConfig(
                model=self.config.llm_model,
                max_tokens=self.config.max_tokens,
                temperature=self.config.temperature,
            ),
            api_key=self.api_key,
        )
        
        # Cache
        if self.config.enable_cache:
            self.cache = RAGCache(
                config=CacheConfig(ttl=self.config.cache_ttl),
            )
        else:
            self.cache = None
    
    def add_documents(
        self,
        sources: Union[str, Path, List[str], List[Path], List[Document]],
        metadata: Optional[Dict[str, Any]] = None,
        batch_size: Optional[int] = None,
    ) -> int:
        """
        Add documents to the RAG pipeline.
        
        Args:
            sources: File paths, URLs, or Document objects
            metadata: Metadata to attach to all documents
            batch_size: Batch size for processing
        
        Returns:
            Number of chunks added
        
        Examples:
            >>> pipeline.add_documents("document.pdf")
            >>> pipeline.add_documents(["doc1.txt", "doc2.txt"])
            >>> pipeline.add_documents(documents, metadata={"source": "manual"})
        """
        batch_size = batch_size or self.config.batch_size
        metadata = metadata or {}
        
        # Normalize input
        if isinstance(sources, (str, Path)):
            sources = [sources]
        
        documents = []
        for source in sources:
            if isinstance(source, Document):
                documents.append(source)
            else:
                doc = self._load_document(source, metadata)
                if doc:
                    documents.append(doc)
        
        if not documents:
            logger.warning("No documents to add")
            return 0
        
        # Process in batches
        total_chunks = 0
        for i in range(0, len(documents), batch_size):
            batch = documents[i:i + batch_size]
            chunks = self._process_documents(batch)
            total_chunks += len(chunks)
        
        logger.info(f"Added {len(documents)} documents, {total_chunks} chunks")
        return total_chunks
    
    def _load_document(
        self,
        source: Union[str, Path],
        metadata: Dict[str, Any],
    ) -> Optional[Document]:
        """Load document from source."""
        source = Path(source)
        
        if not source.exists():
            logger.error(f"Source not found: {source}")
            return None
        
        # Determine document type
        suffix = source.suffix.lower()
        doc_type_map = {
            ".txt": DocumentType.TEXT,
            ".pdf": DocumentType.PDF,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".md": DocumentType.MARKDOWN,
            ".json": DocumentType.JSON,
            ".csv": DocumentType.CSV,
            ".docx": DocumentType.DOCX,
        }
        doc_type = doc_type_map.get(suffix, DocumentType.TEXT)
        
        # Load content based on type
        content = self._extract_content(source, doc_type)
        if not content:
            return None
        
        return Document(
            content=content,
            metadata={**metadata, "filename": source.name},
            doc_type=doc_type,
            source=str(source),
        )
    
    def _extract_content(self, path: Path, doc_type: DocumentType) -> Optional[str]:
        """Extract text content from document."""
        try:
            if doc_type == DocumentType.PDF:
                return self._extract_pdf(path)
            elif doc_type == DocumentType.DOCX:
                return self._extract_docx(path)
            else:
                return path.read_text(encoding="utf-8")
        except Exception as e:
            logger.error(f"Error extracting content from {path}: {e}")
            return None
    
    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        except ImportError:
            logger.warning("PyMuPDF not installed. Reading as text.")
            return path.read_text(errors="ignore")
    
    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX."""
        try:
            import docx
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            logger.warning("python-docx not installed. Reading as text.")
            return path.read_text(errors="ignore")
    
    def _process_documents(self, documents: List[Document]) -> List[Chunk]:
        """Process documents into embedded chunks."""
        all_chunks = []
        
        # Chunk documents
        for doc in documents:
            chunks = self.chunker.chunk(doc)
            all_chunks.extend(chunks)
        
        # Generate embeddings
        texts = [c.content for c in all_chunks]
        embeddings = self.embedder.embed_batch(texts)
        
        for chunk, embedding in zip(all_chunks, embeddings):
            chunk.embedding = embedding
        
        # Add to vector store
        self.vector_store.add(all_chunks)
        
        return all_chunks
    
    def query(
        self,
        question: str,
        filters: Optional[Dict[str, Any]] = None,
        top_k: Optional[int] = None,
        rerank: bool = True,
    ) -> Response:
        """
        Query the RAG pipeline.
        
        Args:
            question: User question
            filters: Metadata filters for retrieval
            top_k: Number of chunks to retrieve
            rerank: Whether to rerank results
        
        Returns:
            Response object with answer and sources
        
        Examples:
            >>> response = pipeline.query("What is machine learning?")
            >>> print(response.answer)
            >>> print(f"Sources: {len(response.sources)}")
        """
        start_time = time.time()
        
        # Check cache
        if self.cache:
            cached = self.cache.get(question)
            if cached:
                cached.cached = True
                return cached
        
        # Create query object
        query = Query(
            text=question,
            filters=filters or {},
            top_k=top_k or self.config.top_k,
            rerank=rerank and self.config.use_reranker,
        )
        
        # Retrieve
        retrieval_start = time.time()
        retrieval_result = self.retriever.retrieve(query)
        retrieval_time = (time.time() - retrieval_start) * 1000
        
        # Rerank
        if query.rerank and self.reranker and retrieval_result.chunks:
            retrieval_result = self.reranker.rerank(query, retrieval_result)
        
        # Generate response
        generation_start = time.time()
        answer = self.generator.generate(query, retrieval_result)
        generation_time = (time.time() - generation_start) * 1000
        
        total_time = (time.time() - start_time) * 1000
        
        # Build response
        response = Response(
            answer=answer,
            sources=retrieval_result.chunks,
            query=query,
            confidence=self._calculate_confidence(retrieval_result),
            retrieval_result=retrieval_result,
            generation_time_ms=generation_time,
            total_time_ms=total_time,
            model_used=self.config.llm_model,
        )
        
        # Update cache
        if self.cache:
            self.cache.set(question, response)
        
        # Update stats
        self._query_count += 1
        self._total_retrieval_time += retrieval_time
        self._total_generation_time += generation_time
        
        return response
    
    def _calculate_confidence(self, result: RetrievalResult) -> float:
        """Calculate answer confidence based on retrieval scores."""
        if not result.scores:
            return 0.0
        
        avg_score = sum(result.scores) / len(result.scores)
        return min(1.0, avg_score)
    
    def search(
        self,
        question: str,
        top_k: int = 10,
        filters: Optional[Dict[str, Any]] = None,
    ) -> RetrievalResult:
        """
        Search without generation (retrieval only).
        
        Args:
            question: Search query
            top_k: Number of results
            filters: Metadata filters
        
        Returns:
            RetrievalResult with matching chunks
        """
        query = Query(text=question, top_k=top_k, filters=filters or {})
        return self.retriever.retrieve(query)
    
    def add_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> int:
        """
        Add raw text directly.
        
        Args:
            text: Text content
            metadata: Optional metadata
        
        Returns:
            Number of chunks added
        """
        doc = Document(content=text, metadata=metadata or {})
        return self.add_documents([doc])
    
    def delete(self, doc_id: str) -> bool:
        """
        Delete a document by ID.
        
        Args:
            doc_id: Document ID to delete
        
        Returns:
            True if deleted successfully
        """
        return self.vector_store.delete(doc_id)
    
    def clear(self) -> None:
        """Clear all documents from the pipeline."""
        self.vector_store.clear()
        if self.cache:
            self.cache.clear()
        logger.info("Pipeline cleared")
    
    def get_stats(self) -> Dict[str, Any]:
        """Get pipeline statistics."""
        return {
            "total_queries": self._query_count,
            "total_documents": self.vector_store.count(),
            "avg_retrieval_time_ms": (
                self._total_retrieval_time / self._query_count
                if self._query_count > 0 else 0
            ),
            "avg_generation_time_ms": (
                self._total_generation_time / self._query_count
                if self._query_count > 0 else 0
            ),
            "cache_size": self.cache.size if self.cache else 0,
            "config": self.config.to_dict(),
        }
    
    def save(self, path: Union[str, Path]) -> None:
        """Save pipeline state to disk."""
        path = Path(path)
        path.mkdir(parents=True, exist_ok=True)
        self.vector_store.save(path / "vector_store")
        logger.info(f"Pipeline saved to {path}")
    
    def load(self, path: Union[str, Path]) -> None:
        """Load pipeline state from disk."""
        path = Path(path)
        self.vector_store.load(path / "vector_store")
        logger.info(f"Pipeline loaded from {path}")
    
    # Component setters for customization
    def set_chunker(self, chunker: Chunker) -> None:
        """Set custom chunker."""
        self.chunker = chunker
    
    def set_embedder(self, embedder: Embedder) -> None:
        """Set custom embedder."""
        self.embedder = embedder
        self.retriever.embedder = embedder
    
    def set_reranker(self, reranker: Reranker) -> None:
        """Set custom reranker."""
        self.reranker = reranker
    
    def set_generator(self, generator: Generator) -> None:
        """Set custom generator."""
        self.generator = generator
