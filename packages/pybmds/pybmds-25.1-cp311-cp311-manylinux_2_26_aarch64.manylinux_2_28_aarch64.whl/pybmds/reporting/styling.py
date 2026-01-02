from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Any, Self

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from pydantic import BaseModel

from ..constants import BMDS_BLANK_VALUE, Dtype
from ..datasets.base import DatasetBase
from ..plotting import close_figure
from ..reporting.footnotes import TableFootnote
from ..utils import citation, ff, four_decimal_formatter

if TYPE_CHECKING:
    from ..models.base import BmdModel
    from ..session import Session


def add_continuous_dataset_footnotes(model: BmdModel, footnotes: TableFootnote):
    if model and model.has_results:
        p_values = model.results.tests.p_values
        footnotes.add_footnote(
            None,
            f"Test 1 Dose Response: {four_decimal_formatter(p_values[0])}",
        )
        footnotes.add_footnote(
            None,
            f"Test 2 Homogeneity of Variance: {four_decimal_formatter(p_values[1])}",
        )
        footnotes.add_footnote(
            None,
            f"Test 3 Variance Model Selection: {four_decimal_formatter(p_values[2])}",
        )


class ReporterStyleGuide(BaseModel):
    portrait_width: float = 6.5
    table: str = "bmdsTbl"
    tbl_header: str = "bmdsTblHeader"
    tbl_body: str = "bmdsTblBody"
    tbl_footnote: str = "bmdsTblFootnote"
    fixed_width: str = "bmdsOutputFile"
    header_1: str = "Heading 1"
    header_2: str = "Heading 2"
    header_3: str = "Heading 3"
    header_4: str = "Heading 4"

    def get_header_style(self, level: int) -> str:
        return getattr(self, f"header_{level}")


class Report(BaseModel):
    document: Any = None
    styles: ReporterStyleGuide

    @classmethod
    def build_default(cls) -> Self:
        fn = Path(__file__).parent / "templates/base.docx"
        # remove empty first paragraph
        doc = Document(str(fn))
        if len(doc.paragraphs) > 0:
            p = doc.paragraphs[0]
            if not p.text and not p.runs:
                el = p._element
                el.getparent().remove(el)
                p._p = p._element = None
        return Report(document=doc, styles=ReporterStyleGuide())


def write_cell(cell, value, style, formatter=ff):
    if value == BMDS_BLANK_VALUE:
        value = "-"
    elif isinstance(value, float):
        value = formatter(value)
    cell.paragraphs[0].text = str(value)
    cell.paragraphs[0].style = style


def set_column_width(column, size_in_inches: float):
    for cell in column.cells:
        cell.width = Inches(size_in_inches)


def add_mpl_figure(document, fig, size_in_inches: float):
    with BytesIO() as f:
        fig.savefig(f)
        document.add_picture(f, width=Inches(size_in_inches))
    fig.clf()
    close_figure(fig)


def write_citation(report: Report, header_level: int):
    styles = report.styles
    header_style = styles.get_header_style(header_level)
    report.document.add_paragraph("Recommended citation", header_style)
    report.document.add_paragraph(
        "Please adapt as appropriate; the citations below capture the package version and "
        "timestamps for easier reproducibility of the analysis."
    )
    report.document.add_paragraph(citation(), styles.fixed_width)


def write_dataset_metadata(report: Report, dataset: DatasetBase):
    if dataset.metadata.name:
        write_setting_p(report, "Name: ", dataset.metadata.name)
    for key, value in dataset.metadata.model_extra.items():
        write_setting_p(report, f"{key.title()}: ", str(value))


def write_dataset_table(
    report: Report, dataset: DatasetBase, long: bool = True, model: BmdModel | None = None
):
    """Write dataset table to word report

    Args:
        report (Report): A report instance
        dataset (DatasetBase): A dataset
        long (bool, optional): Write in long (default) or condensed form.
        model (BmdModel | None, optional): a BmdModel
    """
    styles = report.styles
    footnotes = TableFootnote()

    hdr = styles.tbl_header

    dose_units_text = dataset._get_dose_units_text()
    response_units_text = dataset._get_response_units_text()

    if dataset.dtype is Dtype.CONTINUOUS:
        if long:
            tbl = report.document.add_table(dataset.num_dose_groups + 1, 4, style=styles.table)

            write_cell(tbl.cell(0, 0), "Dose" + dose_units_text, hdr)
            write_cell(tbl.cell(0, 1), "N", hdr)
            write_cell(tbl.cell(0, 2), "Mean" + response_units_text, hdr)
            write_cell(tbl.cell(0, 3), "Std. Dev." + response_units_text, hdr)

            for i, (dose, n, mean, stdev) in enumerate(
                zip(dataset.doses, dataset.ns, dataset.means, dataset.stdevs, strict=True)
            ):
                write_cell(tbl.cell(i + 1, 0), dose, styles.tbl_body)
                write_cell(tbl.cell(i + 1, 1), n, styles.tbl_body)
                write_cell(tbl.cell(i + 1, 2), mean, styles.tbl_body)
                write_cell(tbl.cell(i + 1, 3), stdev, styles.tbl_body)

            width = styles.portrait_width / 4
            for col in tbl.columns:
                set_column_width(col, width)

        else:
            tbl = report.document.add_table(3, dataset.num_dose_groups + 1, style=styles.table)

            write_cell(tbl.cell(0, 0), "Dose" + dose_units_text, hdr)
            write_cell(tbl.cell(1, 0), "N", hdr)
            write_cell(tbl.cell(2, 0), "Mean ± SD" + response_units_text, hdr)

            for i, (dose, n, mean, stdev) in enumerate(
                zip(dataset.doses, dataset.ns, dataset.means, dataset.stdevs, strict=True)
            ):
                write_cell(tbl.cell(0, i + 1), dose, styles.tbl_body)
                write_cell(tbl.cell(1, i + 1), n, styles.tbl_body)
                write_cell(tbl.cell(2, i + 1), f"{mean} ± {stdev}", styles.tbl_body)

            for i, col in enumerate(tbl.columns):
                w = 0.75 if i == 0 else (styles.portrait_width - 0.75) / dataset.num_dose_groups
                set_column_width(col, w)

        add_continuous_dataset_footnotes(model, footnotes)

    elif dataset.dtype is Dtype.DICHOTOMOUS:
        if long:
            tbl = report.document.add_table(dataset.num_dose_groups + 1, 3, style=styles.table)
            write_cell(tbl.cell(0, 0), "Dose" + dose_units_text, hdr)
            write_cell(tbl.cell(0, 1), "N", hdr)
            write_cell(tbl.cell(0, 2), "Incidence", hdr)

            for i, (dose, inc, n) in enumerate(
                zip(dataset.doses, dataset.incidences, dataset.ns, strict=True)
            ):
                write_cell(tbl.cell(i + 1, 0), dose, styles.tbl_body)
                write_cell(tbl.cell(i + 1, 1), n, styles.tbl_body)
                write_cell(tbl.cell(i + 1, 2), inc, styles.tbl_body)

            width = styles.portrait_width / 3
            for col in tbl.columns:
                set_column_width(col, width)

        else:
            tbl = report.document.add_table(2, dataset.num_dose_groups + 1, style=styles.table)

            write_cell(tbl.cell(0, 0), "Dose" + dose_units_text, hdr)
            write_cell(tbl.cell(1, 0), "Affected / Total (%)" + response_units_text, hdr)

            for i, (dose, inc, n) in enumerate(
                zip(dataset.doses, dataset.incidences, dataset.ns, strict=True)
            ):
                frac = inc / float(n)
                write_cell(tbl.cell(0, i + 1), dose, styles.tbl_body)
                write_cell(tbl.cell(1, i + 1), f"{inc}/{n}\n({frac:.1%})", styles.tbl_body)

            for i, col in enumerate(tbl.columns):
                w = 0.75 if i == 0 else (styles.portrait_width - 0.75) / dataset.num_dose_groups
                set_column_width(col, w)

    elif dataset.dtype is Dtype.CONTINUOUS_INDIVIDUAL:
        # aggregate responses by unique doses
        data = {"dose": dataset.individual_doses, "response": dataset.responses}
        df = (
            pd.DataFrame(data, dtype=str)
            .groupby("dose")["response"]
            .agg(list)
            .str.join(", ")
            .reset_index()
        )

        if long:
            tbl = report.document.add_table(
                len(dataset.individual_doses) + 1, 2, style=styles.table
            )
            # add headers
            write_cell(tbl.cell(0, 0), "Dose", hdr)
            write_cell(tbl.cell(0, 1), "Response", hdr)

            # write data
            for i, (dose, response) in enumerate(
                zip(dataset.individual_doses, dataset.responses, strict=True)
            ):
                write_cell(tbl.cell(i + 1, 0), dose, styles.tbl_body)
                write_cell(tbl.cell(i + 1, 1), response, styles.tbl_body)

            for col, width in zip(tbl.columns, [1, styles.portrait_width - 1], strict=True):
                set_column_width(col, width)

        else:
            # create a table
            tbl = report.document.add_table(df.shape[0] + 1, 2, style=styles.table)

            # add headers
            write_cell(tbl.cell(0, 0), "Dose", hdr)
            write_cell(tbl.cell(0, 1), "Response", hdr)

            # write data
            for i, row in df.iterrows():
                write_cell(tbl.cell(i + 1, 0), row.dose, styles.tbl_body)
                write_cell(tbl.cell(i + 1, 1), row.response, styles.tbl_body)

            for col, width in zip(tbl.columns, [1, styles.portrait_width - 1], strict=True):
                set_column_width(col, width)

        add_continuous_dataset_footnotes(model, footnotes)

    elif dataset.dtype is Dtype.NESTED_DICHOTOMOUS:
        tbl = report.document.add_table(len(dataset.doses) + 1, 4, style=styles.table)
        for i, text in enumerate(
            f"Dose{dose_units_text}|Litter N|Incidence|Litter Covariates".split("|")
        ):
            write_cell(tbl.cell(0, i), text, hdr)

        for i, cells in enumerate(
            zip(
                dataset.doses,
                dataset.litter_ns,
                dataset.incidences,
                dataset.litter_covariates,
                strict=True,
            ),
            start=1,
        ):
            for j, text in enumerate(cells):
                write_cell(tbl.cell(i, j), text, styles.tbl_body)

        width = styles.portrait_width / 4
        for col in tbl.columns:
            set_column_width(col, width)

    else:
        raise ValueError("Unknown dtype: {dataset.dtype}")

    # write footnote
    if len(footnotes) > 0:
        footnotes.add_footnote_text(report.document, styles.tbl_footnote)


def write_inputs_table(report: Report, session: Session):
    """Add an input summary table to the document.

    Assumes that all settings are consistent across models in a session; finds the model
    with the maximum multistage/polynomial degree to write inputs.

    Args:
        report (Report): A report object
        session (Session): the current model session

    Raises:
        ValueError: if no models are available in the session
    """
    if len(session.models) == 0:
        raise ValueError("No models available")

    styles = report.styles
    hdr = report.styles.tbl_header
    body = report.styles.tbl_body

    results = session.models[0].results if len(session.models) > 0 else None
    settings = [model.settings for model in session.models]
    content = session.models[0].settings.docx_table_data(settings, results)
    tbl = report.document.add_table(len(content), 2, style=styles.table)
    for idx, (key, value) in enumerate(content.items()):
        write_cell(tbl.cell(idx, 0), key, style=hdr)
        write_cell(tbl.cell(idx, 1), value, style=hdr if idx == 0 else body)


def write_pvalue_header(cell, style):
    # write _P_-Value cell; requires run for italics
    p = cell.paragraphs[0]
    p.style = style
    p.add_run("P").italic = True
    p.add_run("-Value")


def write_base_frequentist_table(report: Report, session: Session):
    if session.dataset.dtype is Dtype.NESTED_DICHOTOMOUS:
        return write_nd_frequentist_table(report, session)
    return write_frequentist_table(report, session)


def write_frequentist_table(report: Report, session: Session):
    styles = report.styles
    hdr = report.styles.tbl_header
    body = report.styles.tbl_body

    footnotes = TableFootnote()
    tbl = report.document.add_table(len(session.models) + 1, 9, style=styles.table)

    write_cell(tbl.cell(0, 0), "Model", style=hdr)
    write_cell(tbl.cell(0, 1), "BMDL", style=hdr)
    write_cell(tbl.cell(0, 2), "BMD", style=hdr)
    write_cell(tbl.cell(0, 3), "BMDU", style=hdr)
    write_pvalue_header(tbl.cell(0, 4), style=hdr)
    write_cell(tbl.cell(0, 5), "AIC", style=hdr)
    write_cell(tbl.cell(0, 6), "Scaled Residual at Control", style=hdr)
    write_cell(tbl.cell(0, 7), "Scaled Residual near BMD", style=hdr)
    write_cell(tbl.cell(0, 8), "Recommendation and Notes", style=hdr)

    # write body
    recommended_index = (
        session.recommender.results.recommended_model_index
        if session.has_recommended_model
        else None
    )
    selected_index = session.selected.model_index
    recommendations = session.recommender.results if session.recommendation_enabled else None
    for idx, model in enumerate(session.models):
        row = idx + 1
        write_cell(tbl.cell(row, 0), model.name(), body)
        if recommended_index == idx:
            footnotes.add_footnote(
                tbl.cell(row, 0).paragraphs[0], "BMDS recommended best fitting model"
            )
        if selected_index == idx:
            footnotes.add_footnote(tbl.cell(row, 0).paragraphs[0], session.selected.notes)
        write_cell(tbl.cell(row, 1), model.results.bmdl, body)
        write_cell(tbl.cell(row, 2), model.results.bmd, body)
        write_cell(tbl.cell(row, 3), model.results.bmdu, body)
        write_cell(tbl.cell(row, 4), model.get_gof_pvalue(), body)
        write_cell(tbl.cell(row, 5), model.results.fit.aic, body)
        write_cell(tbl.cell(row, 6), model.results.gof.residual[0], body)
        write_cell(tbl.cell(row, 7), model.results.gof.roi, body)

        cell = tbl.cell(row, 8)
        if recommendations:
            p = cell.paragraphs[0]
            p.style = body
            run = p.add_run(recommendations.bin_text(idx) + "\n")
            run.bold = True
            p.add_run(recommendations.notes_text(idx))
        else:
            write_cell(tbl.cell(row, 8), "-", body)

    # set column width
    widths = np.array([1.75, 0.8, 0.8, 0.7, 0.7, 0.7, 0.7, 0.7, 1.75])
    widths = widths / (widths.sum() / styles.portrait_width)
    for width, col in zip(widths, tbl.columns, strict=True):
        set_column_width(col, width)

    # write footnote
    if len(footnotes) > 0:
        footnotes.add_footnote_text(report.document, styles.tbl_footnote)


def write_nd_frequentist_table(report: Report, session: Session):
    styles = report.styles
    hdr = report.styles.tbl_header
    body = report.styles.tbl_body

    footnotes = TableFootnote()
    tbl = report.document.add_table(len(session.models) + 1, 7, style=styles.table)

    write_cell(tbl.cell(0, 0), "Model", style=hdr)
    write_cell(tbl.cell(0, 1), "BMDL", style=hdr)
    write_cell(tbl.cell(0, 2), "BMD", style=hdr)
    write_cell(tbl.cell(0, 3), "BMDU", style=hdr)
    write_pvalue_header(tbl.cell(0, 4), style=hdr)
    write_cell(tbl.cell(0, 5), "AIC", style=hdr)
    write_cell(tbl.cell(0, 6), "Recommendation and Notes", style=hdr)

    # write body
    recommended_index = (
        session.recommender.results.recommended_model_index
        if session.has_recommended_model
        else None
    )
    selected_index = session.selected.model_index
    recommendations = session.recommender.results if session.recommendation_enabled else None
    for idx, model in enumerate(session.models):
        row = idx + 1
        write_cell(tbl.cell(row, 0), model.name(), body)
        if recommended_index == idx:
            footnotes.add_footnote(
                tbl.cell(row, 0).paragraphs[0], "BMDS recommended best fitting model"
            )
        if selected_index == idx:
            footnotes.add_footnote(tbl.cell(row, 0).paragraphs[0], session.selected.notes)
        write_cell(tbl.cell(row, 1), model.results.bmdl, body)
        write_cell(tbl.cell(row, 2), model.results.bmd, body)
        write_cell(tbl.cell(row, 3), model.results.bmd, body)
        write_cell(tbl.cell(row, 4), model.results.combined_pvalue, body)
        write_cell(tbl.cell(row, 5), model.results.aic, body)

        cell = tbl.cell(row, 6)
        if recommendations:
            p = cell.paragraphs[0]
            p.style = body
            run = p.add_run(recommendations.bin_text(idx) + "\n")
            run.bold = True
            p.add_run(recommendations.notes_text(idx))
        else:
            write_cell(cell, "-", body)

    # set column width
    widths = np.array([1.5, 0.9, 0.9, 0.9, 0.9, 0.9, 2])
    widths = widths / (widths.sum() / styles.portrait_width)
    for width, col in zip(widths, tbl.columns, strict=True):
        set_column_width(col, width)

    # write footnote
    if len(footnotes) > 0:
        footnotes.add_footnote_text(report.document, styles.tbl_footnote)


def plot_dr(report: Report, session: Session):
    fig = None
    if session.model_average and session.is_bayesian() and session.model_average.has_results:
        fig = session.plot(colorize=False)
    elif session.models[0].has_results:
        fig = session.plot(colorize=True)
    if fig:
        report.document.add_paragraph(add_mpl_figure(report.document, fig, 6))


def write_bayesian_table(report: Report, session: Session):
    styles = report.styles
    report.document.add_paragraph()
    hdr = report.styles.tbl_header
    body = report.styles.tbl_body

    footnotes = TableFootnote()
    tbl = report.document.add_table(len(session.models) + 1, 9, style=styles.table)

    write_cell(tbl.cell(0, 0), "Model", style=hdr)
    write_cell(tbl.cell(0, 1), "Prior Weights", style=hdr)
    write_cell(tbl.cell(0, 2), "Posterior Weights", style=hdr)
    write_cell(tbl.cell(0, 3), "BMDL", style=hdr)
    write_cell(tbl.cell(0, 4), "BMD", style=hdr)
    write_cell(tbl.cell(0, 5), "BMDU", style=hdr)
    write_cell(tbl.cell(0, 6), "Unnormalized Log Posterior Probability", style=hdr)
    write_cell(tbl.cell(0, 8), "Scaled Residual at Control", style=hdr)
    write_cell(tbl.cell(0, 7), "Scaled Residual near BMD", style=hdr)

    ma = session.model_average
    for idx, model in enumerate(session.models, start=1):
        write_cell(tbl.cell(idx, 0), model.name(), body)
        write_cell(tbl.cell(idx, 1), ma.results.priors[idx - 1] if ma else "-", body)
        write_cell(tbl.cell(idx, 2), ma.results.posteriors[idx - 1] if ma else "-", body)
        write_cell(tbl.cell(idx, 3), model.results.bmdl, body)
        write_cell(tbl.cell(idx, 4), model.results.bmd, body)
        write_cell(tbl.cell(idx, 5), model.results.bmdu, body)
        write_cell(tbl.cell(idx, 6), model.results.fit.bic_equiv, body)
        write_cell(tbl.cell(idx, 7), model.results.gof.residual[0], body)
        write_cell(tbl.cell(idx, 8), model.results.gof.roi, body)

    if ma:
        idx = len(tbl.rows)
        tbl.add_row()
        write_cell(tbl.cell(idx, 0), "Model Average", body)
        write_cell(tbl.cell(idx, 1), "-", body)
        write_cell(tbl.cell(idx, 2), "-", body)
        write_cell(tbl.cell(idx, 3), ma.results.bmdl, body)
        write_cell(tbl.cell(idx, 4), ma.results.bmd, body)
        write_cell(tbl.cell(idx, 5), ma.results.bmdu, body)
        write_cell(tbl.cell(idx, 6), "-", body)
        write_cell(tbl.cell(idx, 7), "-", body)
        write_cell(tbl.cell(idx, 8), "-", body)

    # set column width
    widths = np.array([1.1, 0.9, 0.9, 0.9, 0.9, 0.9, 1, 1, 1])
    widths = widths / (widths.sum() / report.styles.portrait_width)
    for width, col in zip(widths, tbl.columns, strict=True):
        set_column_width(col, width)

    # write footnote
    if len(footnotes) > 0:
        footnotes.add_footnote_text(report.document, report.styles.tbl_footnote)


def write_models(report: Report, session: Session, bmd_cdf_table: bool, header_level: int):
    for model in session.models:
        write_model(report, model, bmd_cdf_table, header_level)


def write_model(
    report: Report,
    model: BmdModel,
    bmd_cdf_table: bool,
    header_level: int,
    header_text: str | None = None,
):
    styles = report.styles
    header_style = styles.get_header_style(header_level)
    report.document.add_paragraph(header_text or f"{model.name()} Model", header_style)
    if model.has_results:
        report.document.add_paragraph(add_mpl_figure(report.document, model.plot(), 6))
        if bmd_cdf_table:
            report.document.add_paragraph(add_mpl_figure(report.document, model.cdf_plot(), 6))
    report.document.add_paragraph(model.text(), styles.fixed_width)
    if bmd_cdf_table:
        report.document.add_paragraph("CDF:", styles.tbl_body)
        df_to_table(report, model.cdf())


def write_setting_p(report: Report, title: str, value: str):
    """Write a paragraph with a bolded title string followed by a value."""
    p = report.document.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(value)


def df_to_table(report: Report, df: pd.DataFrame):
    """Quickly generate a word table from a pandas data frame.

    Optimized for speed - see https://github.com/python-openxml/python-docx/issues/174
    """
    hdr = report.styles.tbl_header
    body = report.styles.tbl_body
    n_rows = df.shape[0] + 1
    n_col = df.shape[1]
    tbl = report.document.add_table(n_rows, n_col, style=report.styles.table)
    cells = tbl._cells
    data = df.to_dict("tight", index=False)
    for i, header in enumerate(data["columns"]):
        write_cell(cells[i], header, style=hdr)
    for i, row in enumerate(data["data"]):
        for j, value in enumerate(row):
            write_cell(cells[(i + 1) * n_col + j], value, style=body)
