from docx.document import Document
from docx.text.paragraph import Paragraph


class TableFootnote:
    def __init__(self):
        self.footnotes: dict[str, str] = {}
        self.global_footnotes: list[str] = []

    def __len__(self):
        return len(self.footnotes) + len(self.global_footnotes)

    def _add_footnote_character(self, p, symbol):
        run = p.add_run(symbol)
        run.font.superscript = True

    def add_footnote(self, p: Paragraph | None, text: str):
        """Add a footnote. If a paragraph is provided, a superscript icon added as a run."""
        if p:
            if text not in self.footnotes:
                icon = chr(97 + len(self.footnotes))
                self.footnotes[text] = icon
            self._add_footnote_character(p, self.footnotes[text])
        else:
            if text not in self.global_footnotes:
                self.global_footnotes.append(text)

    def add_footnote_text(self, doc: Document, style: str):
        """Print all footnotes for this instance"""
        for text, char in self.footnotes.items():
            p = doc.add_paragraph("", style=style)
            self._add_footnote_character(p, char)
            p.add_run(f" {text}")
        for text in self.global_footnotes:
            doc.add_paragraph(text, style=style)
